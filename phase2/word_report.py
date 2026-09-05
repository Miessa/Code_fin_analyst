"""Professional Word restitution for the deterministic Phase 2 analysis."""

from datetime import datetime

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


NAVY = "17365D"
BLUE = "D9EAF7"
LIGHT_BLUE = "EAF2F8"
LIGHT_GREY = "F2F2F2"
RED = "FCE4D6"
GREEN = "E2F0D9"


def _shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_text(cell, text, bold=False, color=None, size=8):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text or "—"))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_table(document, headers, rows, widths=None, font_size=8):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    header = table.rows[0]
    header._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for index, title in enumerate(headers):
        _shade(header.cells[index], NAVY)
        _set_cell_text(header.cells[index], title, bold=True, color="FFFFFF", size=font_size)
        if widths:
            header.cells[index].width = Cm(widths[index])
    for row_number, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            _set_cell_text(cells[index], value, size=font_size)
            if row_number % 2:
                _shade(cells[index], LIGHT_GREY)
            if widths:
                cells[index].width = Cm(widths[index])
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def _bullet(document, text, style=None):
    paragraph = document.add_paragraph(style=style or "List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(str(text))


def _heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def _format_indicator(item):
    value = item.get("valeur")
    unit = item.get("unite", "")
    if value is None:
        return "Non calculable"
    if unit.startswith("ratio"):
        return f"{100 * value:.2f} %"
    if unit == "x":
        return f"{value:.2f}x"
    return f"{value:,.2f}".replace(",", " ")


def generer_word(analyse, destination):
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Cm(29.7), Cm(21.0)
    section.top_margin, section.bottom_margin = Cm(1.5), Cm(1.5)
    section.left_margin, section.right_margin = Cm(1.5), Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    for name, size, color in (("Title", 24, NAVY), ("Heading 1", 16, NAVY), ("Heading 2", 12, NAVY)):
        style = document.styles[name]
        style.font.name, style.font.size = "Aptos Display", Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    header = section.header.paragraphs[0]
    header.text = "ARSEL  |  ANALYSE FINANCIÈRE — PHASE 2"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    header.runs[0].font.size = Pt(8)

    footer = section.footer.paragraphs[0]
    footer.text = "Document généré automatiquement à partir des hypothèses validées — revue analyste requise"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(45)
    run = title.add_run("ANALYSE FINANCIÈRE ARSEL")
    run.bold, run.font.size = True, Pt(26)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    subtitle = document.add_paragraph("Analyse déterministe et comparaison aux benchmarks")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    context = analyse.get("metadata", {}).get("context", {})
    project = context.get("project_name") or "Projet analysé"
    info = document.add_paragraph(f"{project}\n{context.get('technology', 'Technologie non renseignée')} — {context.get('geography', 'Localisation non renseignée')}\n{datetime.now():%d/%m/%Y}")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()

    _heading(document, "1. Synthèse exécutive")
    summary_table = document.add_table(rows=1, cols=1)
    summary_table.style = "Table Grid"
    _shade(summary_table.cell(0, 0), BLUE)
    _set_cell_text(summary_table.cell(0, 0), analyse.get("synthese_executive"), size=11)
    document.add_paragraph()

    titles = {
        "technique": "2. Hypothèses techniques", "investissement": "3. Investissement",
        "exploitation": "4. Exploitation", "financement_dette": "5. Financement et dette",
        "fiscalite": "6. Fiscalité", "rentabilite": "7. Rentabilité", "tarif": "8. Tarif",
    }
    for key, title_text in titles.items():
        _heading(document, title_text)
        content = analyse.get("sections", {}).get(key, {})
        for fact in content.get("constats", []):
            _bullet(document, fact)
        if content.get("points_attention"):
            paragraph = document.add_paragraph()
            run = paragraph.add_run("Points d'attention")
            run.bold = True; run.font.color.rgb = RGBColor.from_string("C65911")
            for point in content["points_attention"]:
                _bullet(document, point)

    document.add_page_break()
    _heading(document, "9. Tableau de comparaison des benchmarks")
    rows = [[x.get("cout"), x.get("valeurs_standards"), x.get("valeurs_projets_region"),
             x.get("couts_projet_gds"), x.get("commentaires")]
            for x in analyse.get("tableau_benchmark_detaille", [])]
    _add_table(document,
               ["Coûts", "Valeurs standards", "Projets en développement dans la région", "Projet analysé", "Commentaires"],
               rows, widths=[4.0, 5.0, 5.7, 4.0, 7.0], font_size=7.5)

    sector = analyse.get("comparaison_sectorielle_irena", {})
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Références sectorielles IRENA — base active")
    run.bold = True
    if sector.get("status") != "APPLIED":
        document.add_paragraph("Aucune référence IRENA applicable dans la base active.")
    else:
        sector_rows = [[x.get("metric"), x.get("value"), x.get("unit"), x.get("geography"),
                        x.get("position"), x.get("source_location")]
                       for x in sector.get("comparisons", [])]
        _add_table(document, ["Métrique", "Référence", "Unité", "Géographie", "Position", "Provenance"],
                   sector_rows, widths=[4, 3, 3.5, 3.5, 3, 7], font_size=7)

    _heading(document, "10. Projets comparables approuvés")
    peers = analyse.get("comparaison_projets_pairs", {})
    if peers.get("status") != "APPLIED":
        document.add_paragraph("Comparaison par projets pairs non réalisée ou aucun comparable approuvé.")
    else:
        document.add_paragraph(f"{peers.get('approved_count', 0)} projet(s) approuvé(s) par l'analyste.")
        peer_rows = [[x.get("label"), x.get("project_value"), x.get("p25"), x.get("median"),
                      x.get("p75"), x.get("sample_size"), x.get("position"), x.get("reliability")]
                     for x in peers.get("comparisons", [])]
        _add_table(document, ["Métrique", "Projet", "P25", "Médiane", "P75", "n", "Position", "Fiabilité"],
                   peer_rows, widths=[4.5, 3, 3, 3, 3, 1.5, 3, 2.5], font_size=7)

    _heading(document, "11. Indicateurs financiers dérivés")
    derived = [[x.get("cle"), _format_indicator(x), x.get("unite"), x.get("formule")]
               for x in analyse.get("indicateurs_derives", []) if x.get("calculable")]
    _add_table(document, ["Indicateur", "Valeur", "Unité", "Formule"], derived,
               widths=[5.5, 3.5, 3.5, 10.0], font_size=8)

    for number, title_text, key, fill in (
        (12, "Risques principaux", "risques", RED),
        (13, "Données manquantes importantes", "donnees_manquantes_importantes", LIGHT_GREY),
        (14, "Recommandations", "recommandations", GREEN),
    ):
        _heading(document, f"{number}. {title_text}")
        values = analyse.get(key, [])
        if not values:
            document.add_paragraph("Aucun élément identifié.")
        for value in values:
            table = document.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            _shade(table.cell(0, 0), fill)
            _set_cell_text(table.cell(0, 0), value, size=9)
            document.add_paragraph().paragraph_format.space_after = Pt(1)

    document.core_properties.title = "Analyse financière ARSEL — Phase 2"
    document.core_properties.subject = "Analyse financière et comparaison aux benchmarks"
    document.core_properties.author = "ARSEL Financial Analysis Tool"
    document.save(destination)
    return destination
