from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "outputs" / "rapport_detaille_assets"
TARGET = ROOT / "outputs" / "Sous_section_defis_prealables_phase1_version_corrigee.docx"
BLUE = "17365D"
MID_BLUE = "2E75B6"
PALE_BLUE = "D9EAF7"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Aptos"
    r.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12
    for style_name, size, color in [("Title", 20, BLUE), ("Heading 1", 15, BLUE), ("Heading 2", 12.5, MID_BLUE)]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def paragraphs(doc, *texts):
    for text in texts:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def figure(doc, filename, caption, width=6.45):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ASSETS / filename), width=Inches(width))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("595959")


def add_summary_table(doc):
    rows = [
        ("Localisation de l’unité", "Unité placée au dessus, à droite ou à gauche de la valeur", "Recherche contextuelle verticale et horizontale"),
        ("Devise et échelle", "EUR, XAF, milliers ou montant unitaire peuvent être combinés", "Normalisation conjointe valeur, devise et facteur d’échelle"),
        ("Structure de la cellule", "Une valeur peut être une saisie, une formule ou un total calculé", "Analyse des formules et des dépendances"),
        ("Périmètre économique", "Plusieurs coûts proches ne représentent pas le même contenu", "Règles métier, exclusions et présentation des alternatives"),
        ("Périmètre temporel", "Un taux peut varier selon la période ou le régime applicable", "Conservation du contexte et validation de l’analyste"),
        ("Structure temporelle", "Certaines hypothèses forment une trajectoire plutôt qu’un nombre unique", "Modèle typé pour les séries et les paliers"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, label in enumerate(["Défi", "Risque d’erreur", "Réponse attendue"]):
        set_cell_text(table.rows[0].cells[index], label, True)
        shade(table.rows[0].cells[index], BLUE)
        for run in table.rows[0].cells[index].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_text(cells[index], value, index == 0)
        shade(cells[0], PALE_BLUE)


def build():
    doc = Document()
    configure(doc)
    doc.add_heading("Défis préalables liés à l’extraction des métriques", 0)
    p = doc.add_paragraph("Sous section destinée au chapitre « Flux de la Phase 1 »")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True

    paragraphs(
        doc,
        "L’extraction d’une métrique dans un modèle financier ne consiste pas seulement à retrouver un mot et à lire la cellule placée à sa droite. Dans un classeur non standardisé, la signification d’une valeur dépend à la fois de son libellé, de sa position, de son unité, de sa devise, de son échelle, de sa formule, de la période couverte et du périmètre économique auquel elle se rapporte. Deux cellules proches sur le plan lexical peuvent ainsi représenter des réalités financières très différentes. À l’inverse, la cellule pertinente peut être éloignée du terme employé dans le référentiel ARSEL ou utiliser une expression propre au projet.",
        "Les premières expérimentations ont permis de transformer certaines de ces difficultés générales en problèmes observables. Les captures suivantes proviennent du classeur étudié et montrent pourquoi la Phase 1 doit combiner plusieurs niveaux d’analyse. Elles ne servent pas uniquement à illustrer des erreurs passées. Elles permettent surtout d’expliciter les informations que le système doit réunir avant de pouvoir proposer une métrique avec un niveau de confiance réellement utile à un analyste moins expérimenté.",
        "La liste présentée dans cette sous section est volontairement non exhaustive. Elle développe les difficultés pour lesquelles une capture exploitable a été préparée, mais elle ne couvre pas toutes les contraintes rencontrées pendant l’extraction. D’autres défis déjà identifiés sont récapitulés à la fin de la sous section afin de ne pas donner l’impression que les exemples illustrés définissent à eux seuls l’ensemble du problème.",
    )
    add_summary_table(doc)

    heading(doc, "1. Localiser une unité dont la position n’est pas fixe", 1)
    paragraphs(
        doc,
        "La première difficulté illustrée par les trois captures suivantes concerne la localisation spatiale de l’unité. Dans un modèle financier, l’unité ne se trouve pas systématiquement dans la cellule de la valeur ni dans une colonne réservée à cet usage. Selon la présentation choisie par l’auteur du classeur, elle peut apparaître au dessus d’un ensemble de valeurs, immédiatement à droite d’un montant ou entre le libellé et la valeur, donc à gauche de celle ci. Ces trois configurations appartiennent au même problème d’extraction.",
        "Une recherche limitée à une direction fixe ne peut donc pas être généralisée. Pour chaque valeur candidate, le système doit examiner son voisinage horizontal et vertical, tenir compte de l’alignement des lignes et des colonnes, estimer la portée d’un en tête commun et vérifier que l’unité retrouvée appartient à la famille attendue. La proximité visuelle constitue un indice, mais elle doit être confirmée par la nature de la métrique.",
    )

    heading(doc, "1.1 Unité située au dessus de la valeur", 2)
    paragraphs(
        doc,
        "Dans de nombreux tableaux financiers, l’unité n’est pas répétée à côté de chaque montant. Elle apparaît une seule fois dans l’en tête de la colonne et s’applique à toutes les lignes situées en dessous. La première capture montre la rubrique Uses of funds. Les montants relatifs à l’EPC, aux imprévus, aux frais de développement, au financement et au total général partagent l’indication EUR'000s placée au sommet de la colonne. Une lecture limitée au couple libellé valeur retrouverait bien 1 234 548,8 pour l’EPC ou 2 472 419,7 pour le total, mais elle perdrait l’information selon laquelle ces nombres sont exprimés en milliers d’euros.",
        "Cette omission aurait une conséquence majeure. La valeur 1 234 548,8 ne représente pas environ 1,2 million d’euros, mais environ 1,235 milliard d’euros après application du facteur mille. L’unité doit donc être traitée comme une donnée à part entière et non comme un simple texte décoratif. Le collecteur doit remonter dans la colonne, reconnaître les en têtes communs et déterminer jusqu’où leur portée demeure valide. Il doit également conserver l’expression brute, car la notation EUR'000s peut varier selon les classeurs.",
    )
    figure(doc, "unité_1.png", "Figure 1. Une unité commune placée plusieurs lignes au dessus des montants de la rubrique Uses of funds.", 5.4)

    heading(doc, "1.2 Unité située à droite de la valeur", 2)
    paragraphs(
        doc,
        "La position de l’unité peut changer à l’intérieur d’un même classeur. Dans la deuxième capture, le coût EPC total apparaît sur une ligne qui contient la valeur 1 462 489 et l’unité EUR'000s dans une cellule voisine. Les lignes environnantes utilisent pourtant d’autres familles d’unités : pourcentages, jours, euros par jour et euros par million de mètres cubes. Il serait donc incorrect d’appliquer uniformément l’unité monétaire à toute la zone ou de reprendre automatiquement l’unité de la ligne précédente.",
        "Ce cas impose une recherche locale sensible à la ligne. Le système doit privilégier l’unité explicitement alignée avec la valeur tout en vérifiant sa compatibilité avec la nature attendue de la métrique. Pour un coût de construction, EUR'000s est cohérent; pour une durée de retard, days est cohérent; pour un taux, le symbole pour cent est attendu. La proximité spatiale doit ainsi être combinée à une contrainte sémantique. Sans ce contrôle, une cellule correctement localisée peut être restituée avec l’unité d’une métrique voisine.",
    )
    figure(doc, "Unité_2.png", "Figure 2. Plusieurs familles d’unités coexistent dans une même zone; l’unité EUR'000s doit être rattachée à la ligne du coût EPC.", 6.35)

    heading(doc, "1.3 Unité située à gauche de la valeur", 2)
    paragraphs(
        doc,
        "La troisième capture montre les unités placées entre les libellés et les montants, donc à gauche des valeurs correspondantes. Dans la rubrique consacrée aux coûts fixes d’exploitation, certaines lignes sont exprimées en milliers d’euros et d’autres en milliers de francs CFA. Les libellés sont proches et plusieurs lignes semblent répétées. Une extraction fondée uniquement sur le texte Fixed O&M costs pourrait donc sélectionner une valeur sans identifier l’unité située sur la même ligne, ou lui associer celle de la ligne précédente.",
        "La normalisation doit séparer au moins quatre éléments : la valeur numérique brute, la devise, le facteur d’échelle et la période éventuelle. EUR'000s devient ainsi une valeur en euros accompagnée d’un multiplicateur de mille, tandis que XAF'000s devient une valeur en francs CFA avec le même multiplicateur. La conversion vers une devise commune ne doit intervenir qu’ensuite, avec un taux de change et une date de référence explicites. Cette séquence évite de confondre la lecture de l’unité source avec une conversion financière qui exige ses propres hypothèses.",
    )
    figure(doc, "Unité_3.png", "Figure 3. Coexistence de montants en milliers d’euros et en milliers de francs CFA dans une même rubrique d’O&M.", 6.35)

    heading(doc, "2. Reconnaître qu’une cellule représente un agrégat", 1)
    paragraphs(
        doc,
        "Le libellé d’une cellule ne suffit pas toujours à déterminer son rôle. La capture suivante montre la cellule correspondant au total Uses of funds. Sa formule additionne la dette totale contenue en M56, les postes compris entre M36 et M43, puis les postes compris entre M45 et M48. Cette structure constitue une preuve que la valeur est un total construit à partir de plusieurs composantes et non un simple montant isolé.",
        "L’analyse de la formule apporte donc une information que TF IDF et les embeddings ne peuvent pas produire. Ces méthodes peuvent reconnaître la proximité entre un libellé et la notion d’investissement total, mais elles ne démontrent pas que la cellule couvre effectivement plusieurs catégories de dépenses. Lorsqu’une métrique attend un agrégat, une formule de somme cohérente doit renforcer le candidat. À l’inverse, une cellule portant le mot Total mais faisant seulement référence à un poste élémentaire doit être examinée avec prudence.",
        "Cette preuve structurelle reste toutefois dépendante du périmètre métier. La formule visible inclut les coûts de financement en plus des dépenses de construction. Elle peut donc convenir à l’investissement total, mais pas nécessairement au coût de construction hors coûts financiers. L’analyse correcte exige de comprendre à la fois la structure mathématique de la cellule et la nature des éléments additionnés.",
    )
    figure(doc, "formule_total.png", "Figure 4. La formule de la cellule totale révèle une agrégation de plusieurs groupes de dépenses.", 6.55)

    heading(doc, "3. Départager plusieurs candidats plausibles pour une même métrique", 1)
    paragraphs(
        doc,
        "Pour certaines métriques, plusieurs cellules contiennent des mots fortement compatibles avec la définition recherchée. La capture de la console présente le cas de l’investissement total. Total financing cost, Total Insurance cost, Project Cost, Total cost during development phase et Financing costs Tranche 1 obtiennent tous un score élevé. Pourtant, ces cellules ne décrivent pas le même périmètre. Un coût de financement n’est qu’une composante de l’investissement, tandis que Project Cost semble conceptuellement plus proche du total du projet.",
        "La capture met aussi en évidence une difficulté liée aux unités détectées. Project Cost est associé à years, ce qui est incohérent pour un montant et signale probablement une mauvaise remontée du contexte. De son côté, la proposition principale reçoit Other Financing costs comme unité, alors qu’il s’agit manifestement d’un libellé voisin et non d’une unité monétaire. Un score de confiance élevé ne doit donc pas être interprété comme une garantie de justesse lorsque les différentes preuves se contredisent.",
        "Dans ce type de situation, le système doit présenter les principaux candidats avec leur adresse, leur valeur, leur unité, leur structure et les raisons du classement. Les incompatibilités fortes doivent réduire le score ou provoquer un signalement. Si plusieurs périmètres demeurent plausibles, la décision doit rester ouverte à l’analyste plutôt que d’être masquée derrière une proposition unique. Cette présentation simultanée préfigure le tableau de bord souhaité pour l’interface utilisateur.",
    )
    figure(doc, "Same_possible_for_Same_metric.png", "Figure 5. Plusieurs cellules obtiennent un score élevé pour l’investissement total malgré des périmètres et des unités incompatibles.", 6.55)

    heading(doc, "4. Tenir compte du périmètre temporel et fiscal", 1)
    paragraphs(
        doc,
        "La capture intitulée Valeur correcte, mauvaise période illustre plus précisément un problème de régime fiscal et de temporalité. Deux taux d’impôt sur les sociétés apparaissent comme candidats : 38,5 pour cent pour une période d’exemption et 27,5 pour cent pour une autre situation définie selon le chiffre d’affaires. Chacune de ces valeurs peut être correcte dans son propre contexte, mais aucune ne peut être retenue uniquement parce qu’elle ressemble lexicalement à Corporate tax rate.",
        "Le système doit identifier la condition d’application du taux. Une valeur associée à une exemption, à une phase de construction ou à une période transitoire ne représente pas nécessairement le taux utilisé pendant l’exploitation normale. De même, un taux applicable aux entreprises dont le chiffre d’affaires est inférieur à un seuil ne convient pas automatiquement à un projet qui se situe au dessus de ce seuil. Il faut donc préserver le texte conditionnel entourant la cellule et, lorsque le contexte du projet ne permet pas de conclure, présenter les variantes à l’analyste.",
        "Cet exemple montre que la périodicité ne se limite pas à distinguer un mois d’une année. Elle peut également prendre la forme d’un régime successif dans le temps. Une métrique fiscale peut avoir une valeur pendant une période d’exonération, une autre après l’exonération et éventuellement une troisième selon le niveau d’activité. La représentation finale doit alors pouvoir conserver une série de régimes plutôt que forcer prématurément une valeur scalaire.",
    )
    figure(doc, "Valeur_correcte_mauv_period.png", "Figure 6. Deux taux fiscalement plausibles, mais associés à des conditions et des périodes d’application différentes.", 6.55)

    heading(doc, "5. Distinguer une valeur unique d’une série temporelle", 1)
    paragraphs(
        doc,
        "Certaines hypothèses financières ne sont pas définies par un nombre unique. La dernière capture présente quatre séries de taux d’inflation ou d’indice des prix, répétées sur plusieurs périodes. Le taux applicable au Cameroun passe notamment de 7,22 pour cent à 4,80 pour cent, puis à 3,00 pour cent. Retenir seulement la première cellule produirait une valeur réelle mais incomplète. Calculer une moyenne sans conserver les dates effacerait la trajectoire prévue par le modèle.",
        "La Phase 1 doit donc reconnaître la structure de la donnée avant de la normaliser. Une valeur scalaire contient un nombre et son unité. Une série temporelle contient en plus des périodes ordonnées, éventuellement regroupées en paliers. Le registre validé doit pouvoir conserver ces segments avec leur date de début, leur date de fin, leur valeur et leur source. Cette représentation permettra ensuite à la Phase 2 d’appliquer le bon taux à chaque année plutôt que d’utiliser arbitrairement un seul point de la série.",
        "La distinction est également importante pour le productible, les tarifs, les charges d’exploitation et la disponibilité. Une valeur trimestrielle peut être exacte sans représenter une année complète; un tarif peut changer à une date contractuelle; une disponibilité peut varier entre la montée en charge et l’exploitation stabilisée. Le système doit préférer la structure attendue par la métrique et signaler toute réduction d’une série vers une valeur unique.",
    )
    figure(doc, "Serie_temporelle_Vs_Scalaire.png", "Figure 7. Les hypothèses d’inflation forment des trajectoires par paliers et ne doivent pas être réduites automatiquement à leur première valeur.", 6.55)

    heading(doc, "6. Autres défis identifiés mais non illustrés", 1)
    paragraphs(
        doc,
        "Les captures retenues documentent une partie importante des contraintes, mais plusieurs autres difficultés doivent être prises en compte dans la conception et dans le futur protocole d’évaluation. Leur absence dans les figures ne signifie pas qu’elles sont résolues ou secondaires.",
    )
    other_challenges = [
        "La variabilité des appellations. Une même métrique peut être désignée par un sigle, un synonyme, une traduction ou une expression propre au promoteur. EPC, construction cost et total engineering procurement and construction peuvent renvoyer au même concept avec des degrés de précision différents.",
        "La proximité lexicale entre des périmètres financiers distincts. Total EPC Costs, Total CAPEX, Project Cost et Total Financing Cost partagent certains mots, mais ne comprennent pas les mêmes composantes économiques.",
        "L’éloignement entre le libellé et la valeur. La valeur utile peut se trouver plusieurs colonnes plus loin, dans une autre zone de la feuille ou derrière une référence vers une feuille de synthèse.",
        "La présence de valeurs nulles ou non calculées. Un zéro peut être une hypothèse réelle, une formule inactive, une période non commencée ou une sortie encore indisponible. Pour certaines métriques comme le TRI et le WACC, il constitue un signal d’anomalie qui doit être interprété selon le contexte.",
        "La granularité temporelle inadéquate. Une valeur mensuelle ou trimestrielle peut être correcte dans son contexte sans répondre à une métrique annuelle. Le productible, les charges d’exploitation et la disponibilité sont particulièrement concernés.",
        "Les transformations métier. Une indisponibilité de dix pour cent implique une disponibilité de quatre vingt dix pour cent. La cellule source peut donc être pertinente alors que la valeur finale nécessite une transformation déterministe et traçable.",
        "La confusion entre des durées de nature différente. Une durée de construction, une maturité de dette, une période de grâce et une durée de concession sont toutes exprimées en temps, mais elles ne sont pas interchangeables.",
        "La coexistence de plusieurs variantes légitimes. Le classeur peut contenir plusieurs tranches de dette, plusieurs taux de WHT, plusieurs tarifs ou plusieurs scénarios. Le système doit alors conserver les variantes et leur périmètre au lieu d’en sélectionner arbitrairement une seule.",
        "Les contraintes de performance et de disponibilité des services externes. Une exploration répétée du classeur ou un appel au modèle de langage pour chaque métrique peut rendre l’exécution trop lente et l’exposer aux erreurs de quota ou d’indisponibilité.",
        "La qualité des formules et des valeurs calculées. Le classeur peut contenir des liens externes, des cellules non recalculées ou des formules dont le résultat affiché n’est plus à jour. La valeur visible doit donc être interprétée avec sa provenance technique.",
    ]
    for item in other_challenges:
        p = doc.add_paragraph(item, style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    paragraphs(
        doc,
        "Cette énumération demeure elle même ouverte. L’évaluation sur plusieurs modèles financiers pourra révéler d’autres configurations de feuilles, d’autres conventions d’unités et d’autres ambiguïtés propres aux technologies ou aux montages contractuels. Le référentiel des défis devra donc évoluer avec les cas rencontrés.",
    )

    heading(doc, "7. Conséquences pour la conception de la Phase 1", 1)
    paragraphs(
        doc,
        "Les exemples précédents montrent qu’aucune méthode isolée ne peut résoudre l’ensemble du problème. La recherche lexicale repère efficacement les termes précis; les embeddings couvrent les reformulations; l’analyse du voisinage retrouve les unités; l’examen des formules caractérise les agrégats; les règles métier vérifient la nature, le périmètre et les plages plausibles; enfin, la validation humaine tranche les ambiguïtés qui dépendent du contexte réglementaire ou contractuel.",
        "La précision attendue ne doit pas être confondue avec la production systématique d’une réponse unique. Un système fiable doit aussi savoir reconnaître que plusieurs interprétations demeurent possibles. Dans ce cas, la bonne sortie est une liste courte, structurée et explicable, accompagnée des éléments qui permettent à l’analyste de décider. La confiance doit refléter la convergence des preuves plutôt que le seul score du moteur de recherche.",
        "Ces contraintes justifient l’architecture hybride retenue pour la Phase 1. Elles expliquent également pourquoi la valeur et l’unité doivent pouvoir être corrigées séparément au cours de la validation, pourquoi les séries temporelles nécessitent un modèle typé et pourquoi les transformations métier doivent rester traçables. L’objectif final n’est pas seulement de retrouver une cellule, mais de constituer une métrique exploitable dont la valeur, l’unité, l’échelle, la devise, la période, le périmètre et la provenance sont suffisamment clairs pour alimenter les calculs déterministes de la Phase 2.",
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("ARSEL | Défis préalables de la Phase 1")
    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    build()
