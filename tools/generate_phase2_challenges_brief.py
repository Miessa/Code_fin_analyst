from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "outputs" / "Defis_phase2_synthese_3_pages.docx"
BLUE = "17365D"
MID_BLUE = "2E75B6"
PALE = "D9EAF7"


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    item = props.find(qn("w:shd"))
    if item is None:
        item = OxmlElement("w:shd")
        props.append(item)
    item.set(qn("w:fill"), fill)


def cell_text(cell, value, bold=False, white=False):
    cell.text = ""
    run = cell.paragraphs[0].add_run(value)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(8.4)
    if white:
        run.font.color.rgb = RGBColor(255, 255, 255)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.3)
    normal.paragraph_format.space_after = Pt(3.5)
    normal.paragraph_format.line_spacing = 1.02
    for name, size in [("Title", 18), ("Heading 1", 12.5), ("Heading 2", 10.5)]:
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE if name != "Heading 2" else MID_BLUE)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(3)

    title = doc.add_heading("Principaux défis de la Phase 2", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Normalisation, analyse financière et comparaison avec les benchmarks")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    paragraph(doc, "La Phase 2 reçoit les métriques extraites puis validées au terme de la Phase 1. Son rôle est de les transformer en indicateurs financiers cohérents, de sélectionner des références comparables et de produire une analyse professionnelle. Le défi ne consiste donc plus à retrouver une cellule dans Excel, mais à vérifier que les valeurs peuvent réellement être calculées, rapprochées et interprétées sans perdre leur unité, leur période, leur périmètre ni leur provenance. Les difficultés présentées ci dessous sont les principales contraintes déjà identifiées; cette liste demeure non exhaustive et devra être complétée à mesure que de nouveaux modèles financiers seront étudiés.")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, value in enumerate(["Enjeu", "Risque principal", "Réponse attendue"]):
        cell_text(table.rows[0].cells[i], value, True, True)
        shade(table.rows[0].cells[i], BLUE)
    rows = [
        ("Normalisation", "Comparer des formats, devises, échelles ou périodes incompatibles", "Objet structuré valeur, unité, devise, échelle, année et périmètre"),
        ("Calculs", "Produire des ratios numériquement exacts mais économiquement faux", "Formules déterministes, contrôles de cohérence et gestion des métriques complexes"),
        ("Comparables", "Mélanger des technologies, projets ou statistiques non comparables", "Filtres explicites, revue humaine et séparation projets secteur"),
        ("Statistiques", "Surinterpréter un petit échantillon ou des données manquantes", "Effectifs par métrique, médiane, quartiles, valeurs extrêmes et réserves"),
        ("Analyse", "Transformer un écart en conclusion causale non démontrée", "Constats déterministes séparés des hypothèses d’interprétation"),
        ("Restitution", "Perdre la provenance ou laisser le LLM modifier les faits", "Dossier de faits, citations, manifeste et validation de l’analyste"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value, i == 0)
        shade(cells[0], PALE)

    doc.add_heading("1. Normaliser les valeurs avant toute comparaison", 1)
    paragraph(doc, "Les valeurs validées peuvent encore être représentées sous des formes hétérogènes : 13,5791 %, 0,135791, 10 % p.a., 18 699,161 kEUR ou 18 ans T1 et 21 ans T2. La Phase 2 doit convertir ces écritures en objets structurés sans supprimer l’information portée par la forme originale. Une valeur financière complète associe au minimum un nombre, une famille d’unité, une devise, un facteur d’échelle, une période de référence et un périmètre économique.")
    paragraph(doc, "L’uniformisation monétaire est particulièrement sensible. Un montant en kEUR doit être multiplié par mille avant comparaison avec des euros, puis éventuellement converti vers une autre devise avec un taux, une date et une source conservés. Les montants d’années différentes peuvent également devoir être corrigés par un indice de prix approprié. Une conversion silencieuse au taux courant ou une confusion entre EUR et EUR'000s peut créer un écart artificiel bien supérieur à celui que l’analyse cherche à mesurer.")
    paragraph(doc, "Le périmètre doit être normalisé avec autant de soin que l’unité. Un coût EPC, un CAPEX incluant les taxes et un investissement total comprenant les frais financiers ne sont pas interchangeables. La Phase 2 doit soit reconstruire un périmètre commun, soit refuser le calcul comparatif et expliquer la différence restante.")

    doc.add_heading("2. Calculer des indicateurs économiquement cohérents", 1)
    paragraph(doc, "Les indicateurs dérivés doivent être calculés par des formules déterministes et testées. L’investissement par MW exige un investissement et une puissance exprimés dans des unités compatibles. Le facteur de charge repose sur un productible annuel et non trimestriel : facteur de charge = productible annuel divisé par puissance multipliée par 8 760 heures. La part de dette exige une dette et un investissement total appartenant au même périmètre. Un calcul peut être mathématiquement correct tout en étant économiquement faux si l’une de ces conditions n’est pas satisfaite.")
    paragraph(doc, "Certaines métriques ne sont pas scalaires. L’inflation peut évoluer par paliers, la dette peut comprendre plusieurs tranches, et les tarifs ou la disponibilité peuvent former des séries temporelles. Le système doit préserver ces structures plutôt que les réduire automatiquement à une moyenne. Il doit aussi gérer les valeurs nulles, manquantes ou suspectes : zéro peut être une hypothèse réelle, mais aussi une formule inactive ou une sortie non calculée.")

    doc.add_heading("3. Construire un groupe de comparaison réellement pertinent", 1)
    paragraph(doc, "La présence d’un projet dans DuckDB ne suffit pas à en faire un comparable. La technologie doit généralement constituer un filtre strict, puis la sélection doit considérer la géographie, la capacité, la période, le stade de développement, la durée contractuelle, la structure de financement et la disponibilité des métriques. Des filtres trop larges produisent un échantillon abondant mais peu pertinent; des filtres trop stricts peuvent ne laisser que quelques observations. L’analyste doit pouvoir voir les candidats, comprendre les critères et approuver la sélection.")
    paragraph(doc, "Les projets individuels de la Banque mondiale et les statistiques sectorielles de l’IRENA doivent suivre deux voies distinctes. Les premiers constituent un groupe de pairs; les secondes positionnent le projet par rapport à une distribution ou une moyenne sectorielle. Les mélanger dans un même échantillon reviendrait à traiter une statistique agrégée comme un projet réel. Les doublons entre éditions doivent également être éliminés afin qu’un même projet ne pèse pas plusieurs fois.")

    doc.add_heading("4. Produire des statistiques prudentes avec des données imparfaites", 1)
    paragraph(doc, "Après filtrage, l’échantillon peut être réduit et incomplet. Le nombre d’observations doit être affiché séparément pour chaque métrique, car douze projets peuvent renseigner l’investissement alors que trois seulement publient un TRI des fonds propres. La médiane, les quartiles, le minimum et le maximum sont souvent plus informatifs qu’une moyenne isolée, mais ils ne doivent pas créer une impression de précision excessive lorsque l’effectif est faible.")
    paragraph(doc, "Les valeurs extrêmes ne doivent être ni acceptées aveuglément ni supprimées automatiquement. Elles peuvent révéler une erreur d’unité, une différence de périmètre, un contexte géographique particulier ou un projet réellement atypique. Toute exclusion doit être justifiée et enregistrée. La qualité de chaque observation dépend également de sa source, de son année, de son statut de validation et du snapshot dont elle provient.")

    doc.add_heading("5. Passer du constat numérique à une analyse financière défendable", 1)
    paragraph(doc, "Le moteur doit d’abord produire des constats déterministes : position du projet par rapport à la médiane, écart au quartile supérieur, cohérence entre TRI et WACC, marge offerte par le DSCR ou poids de la dette. Il doit ensuite distinguer ces constats des explications possibles. Un coût par MW supérieur au benchmark ne démontre pas une inefficacité; il peut résulter des travaux civils, du raccordement, des taxes, du calendrier ou du périmètre retenu. Les causes proposées doivent être présentées comme des hypothèses à examiner lorsqu’elles ne sont pas établies par les données.")
    paragraph(doc, "Cette séparation est essentielle pour l’intégration future du LLM. Le modèle devra recevoir un dossier de faits contenant les valeurs normalisées, les indicateurs, les benchmarks approuvés, les écarts, les sources, les réserves et les données manquantes. Il pourra ajouter du contexte et approfondir l’interprétation, mais il ne devra ni recalculer les chiffres ni inventer une référence. Les commentaires déterministes resteront disponibles indépendamment de son intervention.")

    doc.add_heading("6. Garantir une restitution lisible, traçable et généralisable", 1)
    paragraph(doc, "Le rapport final doit servir deux niveaux de lecture. Le résumé exécutif présente les conclusions essentielles; les tableaux et annexes permettent à l’analyste de vérifier les valeurs, les échantillons et les sources. Chaque conclusion doit pouvoir être reliée à un indicateur, puis aux métriques validées et aux cellules Excel. Chaque comparaison doit conduire aux observations de DuckDB, à leur source et au snapshot immuable utilisé. Le manifeste d’analyse doit identifier les entrées et sorties exactes.")
    paragraph(doc, "Enfin, la fiabilité de la Phase 2 ne pourra pas être conclue sur le seul modèle Kikot. Elle devra être évaluée sur plusieurs projets, technologies, devises, tailles, structures contractuelles et niveaux de complétude. Le principal défi est de conserver les mêmes règles de calcul et de traçabilité lorsque la forme des données change. La Phase 2 sera réellement robuste lorsqu’elle saura produire une analyse utile, mais aussi reconnaître et expliquer les situations dans lesquelles une comparaison ne peut pas être défendue.")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("ARSEL | Principaux défis de la Phase 2")
    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    build()
