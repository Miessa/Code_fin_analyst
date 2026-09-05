from __future__ import annotations

import json
import re
import textwrap
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
ASSETS = OUT / "rapport_detaille_assets"
SOURCE_MD = OUT / "Rapport_avancement_ARSEL_30_aout_2026.md"
TARGET = OUT / "Rapport_avancement_detaille_ARSEL_31_aout_2026_architecture_cible.docx"

BLUE = "17365D"
MID_BLUE = "2F75B5"
LIGHT_BLUE = "D9EAF7"
PALE = "F3F6F9"
ORANGE = "ED7D31"
GREEN = "70AD47"
GREY = "666666"


def clean(text: str) -> str:
    text = text.replace(" - ", " : ").replace("–", "à").replace("—", ":")
    return text.replace("-", " ").strip()


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, value: str, bold=False, color=None, size=9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(clean(value))
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def footer(section) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Rapport d’avancement détaillé ARSEL   |   31 août 2026   |   ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)
    footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [("Title", 28, BLUE), ("Heading 1", 20, BLUE), ("Heading 2", 15, MID_BLUE), ("Heading 3", 12, BLUE)]:
        st = doc.styles[name]
        st.font.name = "Aptos Display"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True

    if "Légende" not in [s.name for s in doc.styles]:
        st = doc.styles.add_style("Légende", WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = "Aptos"
        st.font.size = Pt(8.5)
        st.font.italic = True
        st.font.color.rgb = RGBColor.from_string(GREY)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def title(doc: Document, text: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1"
    p.add_run(clean(text))
    if subtitle:
        q = doc.add_paragraph(clean(subtitle))
        q.style = "Subtitle"


def page(doc: Document, heading: str, paragraphs: list[str], subtitle: str | None = None) -> None:
    title(doc, heading, subtitle)
    for raw in paragraphs:
        p = doc.add_paragraph(clean(raw))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def callout(doc: Document, heading: str, body: str, fill=LIGHT_BLUE) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 3"
    p.add_run(clean(heading))
    q = doc.add_paragraph(clean(body))
    q.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def box(ax, x, y, w, h, text, color=LIGHT_BLUE, edge=MID_BLUE, fontsize=10, linestyle="-"):
    rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.025", fc="#" + color, ec="#" + edge, lw=1.6, linestyle=linestyle)
    ax.add_patch(rect)
    ax.text(x + w / 2, y + h / 2, clean(text), ha="center", va="center", fontsize=fontsize, wrap=True)


def arrow(ax, x1, y1, x2, y2):
    ax.annotate("", xy=(x2, y2), xytext=(x1, y1), arrowprops=dict(arrowstyle="->", lw=1.8, color="#44546A"))


def architecture_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(12, 6.5))
    ax.set_xlim(0, 12); ax.set_ylim(0, 7); ax.axis("off")
    box(ax, .3, 2.8, 1.6, 1.2, "Classeur Excel\ndu promoteur", "FFF2CC", "BF9000")
    box(ax, 2.3, 2.6, 2.2, 1.6, "Phase 1\nExtraction hybride\net validation", LIGHT_BLUE, MID_BLUE)
    box(ax, 5.0, 2.8, 1.8, 1.2, "Registre des\nhypothèses\nvalidées", "E2F0D9", GREEN)
    box(ax, 7.25, 4.7, 2.0, 1.3, "Benchmark bank\nWorld Bank et IRENA\nDuckDB", "E4DFEC", "7030A0")
    box(ax, 7.25, 1.8, 2.0, 1.3, "Phase 2\nCalculs et\ncomparaisons", LIGHT_BLUE, MID_BLUE)
    box(ax, 10.0, 1.8, 1.65, 1.3, "Rapports\nJSON, Markdown\net Word", "E2F0D9", GREEN)
    box(ax, 7.25, .15, 2.0, .9, "Analyste\nvalidation et décision", "FFF2CC", "BF9000")
    arrow(ax, 1.9, 3.4, 2.3, 3.4); arrow(ax, 4.5, 3.4, 5.0, 3.4)
    arrow(ax, 6.8, 3.4, 7.25, 2.55); arrow(ax, 8.25, 4.7, 8.25, 3.1)
    arrow(ax, 9.25, 2.45, 10.0, 2.45); arrow(ax, 8.25, 1.05, 8.25, 1.8)
    ax.set_title("Architecture fonctionnelle actuelle", fontsize=18, color="#" + BLUE, weight="bold")
    fig.tight_layout(); fig.savefig(path, dpi=180, bbox_inches="tight"); plt.close(fig)


def detailed_architecture_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(17, 11))
    ax.set_xlim(0, 17); ax.set_ylim(0, 11); ax.axis("off")
    ax.text(8.5, 10.65, "Architecture cible détaillée du système", ha="center", fontsize=21, color="#" + BLUE, weight="bold")
    ax.text(.35, 10.15, "LÉGENDE", fontsize=9.5, weight="bold", color="#" + BLUE)
    box(ax, 1.25, 9.92, 1.55, .38, "Déjà conçu", LIGHT_BLUE, MID_BLUE, 8)
    box(ax, 3.05, 9.92, 1.8, .38, "Prévu ou à consolider", "FFF2CC", ORANGE, 8, "--")
    box(ax, 5.1, 9.92, 1.65, .38, "Intervention LLM", "FCE4D6", ORANGE, 8)

    ax.text(.35, 9.38, "ENTRÉES", fontsize=10.5, weight="bold", color="#" + BLUE)
    box(ax, .3, 8.25, 1.8, .72, "Classeur Excel\nnon standardisé", "FFF2CC", "BF9000", 8.3)
    box(ax, .3, 7.25, 1.8, .72, "Référentiel ARSEL\nconcepts et définitions", "FFF2CC", "BF9000", 8.3)
    box(ax, .3, 6.25, 1.8, .72, "Contexte du projet\ntechnologie, pays, contrats", "FFF2CC", "BF9000", 8.3)

    ax.text(2.5, 9.38, "ÉTAPE 0   COMPRÉHENSION ASSISTÉE", fontsize=10.5, weight="bold", color="#" + BLUE)
    box(ax, 2.5, 8.15, 2.15, .88, "0A. Définitions déterministes\ndes métriques ARSEL", LIGHT_BLUE, MID_BLUE, 8.5)
    box(ax, 5.0, 8.15, 2.15, .88, "0B. Détection des termes\npropres au dossier", "FFF2CC", ORANGE, 8.5, "--")
    box(ax, 7.5, 8.15, 2.3, .88, "0C. LLM : glossaire dynamique\nEPC, WHT et expressions inconnues", "FCE4D6", ORANGE, 8.5, "--")
    box(ax, 10.15, 8.15, 2.1, .88, "0D. Explications adaptées\nà l’analyste", "FFF2CC", ORANGE, 8.5, "--")
    arrow(ax, 2.1, 7.62, 2.5, 8.45); arrow(ax, 4.65, 8.59, 5.0, 8.59); arrow(ax, 7.15, 8.59, 7.5, 8.59); arrow(ax, 9.8, 8.59, 10.15, 8.59)

    ax.text(2.5, 7.55, "PHASE 1   EXTRACTION ET VALIDATION", fontsize=10.5, weight="bold", color="#" + BLUE)
    phase1 = [
        ("1. Collecte\nvaleurs et contexte", 2.5, LIGHT_BLUE, MID_BLUE, "-"),
        ("2. TF IDF\net embeddings", 4.55, LIGHT_BLUE, MID_BLUE, "-"),
        ("3. Fusion et preuves\nstructure, unité, métier", 6.6, LIGHT_BLUE, MID_BLUE, "-"),
        ("4. LLM si ambiguïté\nsur shortlist fermée", 8.85, "FCE4D6", ORANGE, "-"),
        ("5. Validation analyste\nvaleur et unité", 11.1, LIGHT_BLUE, MID_BLUE, "-"),
        ("6. Métriques validées\nhypotheses_validees.json", 13.35, "E2F0D9", GREEN, "-"),
    ]
    for label, x, fc, ec, ls in phase1: box(ax, x, 6.42, 1.85, .78, label, fc, ec, 8.2, ls)
    for i in range(len(phase1)-1): arrow(ax, phase1[i][1]+1.85, 6.81, phase1[i+1][1], 6.81)
    arrow(ax, 2.1, 8.58, 2.5, 6.92); arrow(ax, 2.1, 7.58, 4.55, 6.92)

    ax.text(2.5, 5.72, "PHASE 2   CALCULS, COMPARAISONS ET ANALYSE", fontsize=10.5, weight="bold", color="#" + BLUE)
    phase2 = [
        ("12. Normalisation\net indicateurs dérivés",2.5,LIGHT_BLUE,MID_BLUE,"-"),
        ("13. Comparables et\nbenchmarks sectoriels",4.75,LIGHT_BLUE,MID_BLUE,"-"),
        ("14. Constats financiers\ndéterministes",7.0,LIGHT_BLUE,MID_BLUE,"-"),
        ("15. Dossier de faits\nécarts, risques, sources",9.25,"FFF2CC",ORANGE,"--"),
        ("16. LLM : contexte et\ninterprétation approfondie",11.5,"FCE4D6",ORANGE,"--"),
        ("17. Revue analyste et\nrapport final",14.05,"E2F0D9",GREEN,"--"),
    ]
    for label,x,fc,ec,ls in phase2: box(ax,x,4.42,1.9 if x<14 else 1.75,.82,label,fc,ec,8.1,ls)
    for i in range(len(phase2)-1):
        width=1.9 if phase2[i][1]<14 else 1.75; arrow(ax,phase2[i][1]+width,4.83,phase2[i+1][1],4.83)

    # Les deux flux verticaux rejoignent la Phase 2 par des couloirs libres,
    # sans traverser les blocs de traitement.
    ax.plot([14.275, 14.275, 3.45], [6.42, 5.48, 5.48], color="#44546A", lw=1.35)
    arrow(ax, 3.45, 5.48, 3.45, 5.24)
    ax.plot([2.1, 2.27, 2.27], [6.62, 5.48, 4.83], color="#44546A", lw=1.35)
    arrow(ax, 2.27, 4.83, 2.5, 4.83)

    ax.text(.35, 3.82, "BANQUE DE BENCHMARKS", fontsize=10.5, weight="bold", color="#" + BLUE)
    box(ax, .3, 2.92, 1.8, .68, "World Bank PPI\nprojets individuels", "E4DFEC", "7030A0", 8.2)
    box(ax, .3, 2.05, 1.8, .68, "IRENA tabulaire\nstatistiques sectorielles", "E4DFEC", "7030A0", 8.2)
    bank = [("7. Détection et\nsnapshot SHA 256",2.5),("8. Adaptateurs et\nnormalisation",4.55),("9. Staging et\ncontrôles qualité",6.6),("10. Revue et\npromotion humaine",8.65),("11. DuckDB active\nprojets et observations",10.7)]
    for label,x in bank: box(ax,x,2.42,1.75,.78,label,"E4DFEC","7030A0",8.1)
    for i in range(len(bank)-1): arrow(ax,bank[i][1]+1.75,2.81,bank[i+1][1],2.81)
    arrow(ax,2.1,3.26,2.5,2.98); arrow(ax,2.1,2.39,2.5,2.63)

    # La base active alimente les comparables par une liaison orthogonale
    # placée dans l’espace séparant la banque de la Phase 2.
    ax.plot([12.45, 12.45, 5.70], [2.81, 3.78, 3.78], color="#44546A", lw=1.35)
    arrow(ax, 5.70, 3.78, 5.70, 4.42)

    ax.text(2.5, 1.62, "MÉCANISMES TRANSVERSAUX DÉJÀ CONÇUS", fontsize=10.5, weight="bold", color="#" + BLUE)
    cross=[("Provenance\net sources",2.5),("Instrumentation\ndes appels LLM",4.55),("Tests et\nrègles métier",6.6),("Checksums et\nmanifestes",8.65),("Journal des décisions\nde l’analyste",10.7),("Confidentialité et\nreproductibilité",12.75)]
    for label,x in cross: box(ax,x,.62,1.75,.68,label,PALE,"7F8C8D",8.1)
    ax.text(8.5,.14,"Le code conserve les calculs et la conduite du pipeline. Le LLM intervient à trois points fonctionnels bornés. L’analyste conserve la validation et la décision.",ha="center",fontsize=9.5,color="#44546A")
    fig.tight_layout(); fig.savefig(path, dpi=190, bbox_inches="tight"); plt.close(fig)


def retrieval_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(12, 6.5))
    ax.set_xlim(0, 12); ax.set_ylim(0, 7); ax.axis("off")
    labels = [
        ("Référentiel ARSEL\nsynonymes et attentes", .3, 4.7, "FFF2CC", "BF9000"),
        ("Collecte des libellés\net cellules candidates", 2.5, 4.7, LIGHT_BLUE, MID_BLUE),
        ("Recherche lexicale\nTF IDF", 5.0, 5.3, LIGHT_BLUE, MID_BLUE),
        ("Recherche sémantique\nembeddings locaux", 5.0, 3.8, "E4DFEC", "7030A0"),
        ("Règles structurelles\nformules, agrégats, unités", 7.5, 4.7, "E2F0D9", GREEN),
        ("Fusion des rangs\net score de confiance", 9.8, 4.7, LIGHT_BLUE, MID_BLUE),
        ("LLM uniquement\nsi ambiguïté réelle", 7.5, 2.2, "FCE4D6", ORANGE),
        ("Proposition et\nalternatives explicables", 9.8, 2.2, "E2F0D9", GREEN),
    ]
    for text, x, y, fc, ec in labels: box(ax, x, y, 1.9, 1.0, text, fc, ec, 9)
    arrow(ax, 2.2, 5.2, 2.5, 5.2); arrow(ax, 4.4, 5.2, 5.0, 5.8); arrow(ax, 4.4, 5.0, 5.0, 4.3)
    arrow(ax, 6.9, 5.8, 7.5, 5.2); arrow(ax, 6.9, 4.3, 7.5, 5.0); arrow(ax, 9.4, 5.2, 9.8, 5.2)
    arrow(ax, 10.75, 4.7, 8.45, 3.2); arrow(ax, 9.4, 2.7, 9.8, 2.7)
    ax.set_title("Chaîne hybride de recherche des métriques", fontsize=18, color="#" + BLUE, weight="bold")
    fig.tight_layout(); fig.savefig(path, dpi=180, bbox_inches="tight"); plt.close(fig)


def bank_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(12, 6.5))
    ax.set_xlim(0, 12); ax.set_ylim(0, 7); ax.axis("off")
    box(ax, .3, 4.8, 1.8, 1.0, "World Bank PPI\nprojets individuels", LIGHT_BLUE, MID_BLUE)
    box(ax, .3, 2.8, 1.8, 1.0, "IRENA\nstatistiques sectorielles", "E4DFEC", "7030A0")
    box(ax, 2.7, 3.8, 1.7, 1.2, "Téléchargement\ntemporaire", "FFF2CC", "BF9000")
    box(ax, 5.0, 3.8, 1.8, 1.2, "Snapshot immuable\nadressé par SHA 256", "E2F0D9", GREEN)
    box(ax, 7.4, 4.8, 1.7, 1.0, "Normalisation\net contrôles", LIGHT_BLUE, MID_BLUE)
    box(ax, 7.4, 2.8, 1.7, 1.0, "Staging\nDuckDB", "FCE4D6", ORANGE)
    box(ax, 9.8, 3.8, 1.8, 1.2, "Promotion humaine\nvers la base active", "E2F0D9", GREEN)
    arrow(ax, 2.1, 5.3, 2.7, 4.6); arrow(ax, 2.1, 3.3, 2.7, 4.2); arrow(ax, 4.4, 4.4, 5.0, 4.4)
    arrow(ax, 6.8, 4.4, 7.4, 5.3); arrow(ax, 8.25, 4.8, 8.25, 3.8); arrow(ax, 9.1, 3.3, 9.8, 4.2)
    ax.set_title("Ingestion, gouvernance et promotion des benchmarks", fontsize=18, color="#" + BLUE, weight="bold")
    fig.tight_layout(); fig.savefig(path, dpi=180, bbox_inches="tight"); plt.close(fig)


def provenance_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(12, 6.5))
    ax.set_xlim(0, 12); ax.set_ylim(0, 7); ax.axis("off")
    xs = [.4, 2.6, 4.8, 7.0, 9.3]
    labs = ["Cellule et formule\nd’origine", "Candidat classé\navec preuves", "Décision et correction\nde l’analyste", "Calcul déterministe\net comparaison", "Conclusion du rapport\net manifeste"]
    cols = ["FFF2CC", LIGHT_BLUE, "E4DFEC", LIGHT_BLUE, "E2F0D9"]
    edges = ["BF9000", MID_BLUE, "7030A0", MID_BLUE, GREEN]
    for x, lab, fc, ec in zip(xs, labs, cols, edges): box(ax, x, 3.0, 1.8, 1.3, lab, fc, ec, 9)
    for i in range(4): arrow(ax, xs[i] + 1.8, 3.65, xs[i+1], 3.65)
    ax.text(6, 1.7, "Chaque transformation conserve sa source, son unité, son statut et sa justification", ha="center", fontsize=12, color="#" + BLUE)
    ax.set_title("Chaîne de provenance et de justification", fontsize=18, color="#" + BLUE, weight="bold")
    fig.tight_layout(); fig.savefig(path, dpi=180, bbox_inches="tight"); plt.close(fig)


def console_figure(path: Path, kind: str) -> None:
    if kind == "slow":
        lines = [
            "EXÉCUTION RÉELLE   branche origin/main",
            "ÉTAPE 1   Extraction  présélection  score  rerank",
            "Collecte des libellés… 4192 libellés.  2.2 s",
            "Gemini : disponible",
            "Ouverture du modèle : 0.5 s",
            "",
            "  Routing LLM : OUI",
            "  Gemini erreur temporaire [json] : 503 UNAVAILABLE",
            "  Gemini erreur temporaire [tool_calling:cout_construction] : 503 UNAVAILABLE",
            "  Gemini erreur temporaire [json] : 503 UNAVAILABLE",
            "cout_construction  →  à explorer",
            "  Aucun candidat supplémentaire trouvé  exploration requise",
            "",
            "  Routing LLM : OUI",
            "  Gemini erreur temporaire [tool_calling:investissement_total] : 503 UNAVAILABLE",
            "investissement_total  →  à explorer",
        ]
    else:
        lines = [
            "EXÉCUTION RÉELLE   branche origin/feature/architecture-hybride",
            "ÉTAPE 1   Extraction  présélection  score  rerank",
            "Collecte des libellés… 5606 libellés.  2.7 s",
            "Ouverture du modèle : 80.8 s",
            "",
            "  Routing LLM : OUI",
            "  Gemini quota/rate limit [json]",
            "  Sélection : AMBIGUOUS / INVALID_RESPONSE",
            "  shortlist déterministe conservée pour validation",
            "cout_construction  →  proposé  [100 %]  InpC!F515 = 1 462 489,23  (7.5 s)",
            "",
            "                     […] extraits non contigus […]",
            "duree_dette  →  proposé  [79 %]  InpC!F687 = 18  (109.1 s)",
            "is_taux      →  proposé  [100 %] InpC!F1358 = 0.385  (113.9 s)",
            "wht          →  proposé  [100 %] InpC!F1546 = 0  (102.5 s)",
            "",
            "Appels LLM : 22  |  succès pipeline=22  |  échecs instrumentés=0",
        ]
    fig, ax = plt.subplots(figsize=(12, 6.5), facecolor="#111827")
    ax.set_facecolor("#111827"); ax.axis("off")
    ax.text(.035, .95, "●  ●  ●", color="#F87171", fontsize=16, va="top", transform=ax.transAxes)
    ax.text(.04, .85, "\n".join(lines), family="DejaVu Sans Mono", fontsize=10.5, color="#E5E7EB", va="top", transform=ax.transAxes, linespacing=1.28)
    fig.tight_layout(); fig.savefig(path, dpi=180, facecolor=fig.get_facecolor(), bbox_inches="tight"); plt.close(fig)


def add_figure_page(doc: Document, heading: str, path: Path, caption: str, note: str) -> None:
    title(doc, heading)
    doc.add_picture(str(path), width=Inches(6.8))
    p = doc.add_paragraph(clean(caption)); p.style = "Légende"
    callout(doc, "Lecture", note)
    doc.add_page_break()


def add_capture_analysis_page(
    doc: Document,
    heading: str,
    path: Path,
    caption: str,
    observation: str,
    correction: str,
    reserve: str | None = None,
) -> None:
    """Insert an original execution screenshot and interpret it in place."""
    title(doc, heading)
    doc.add_picture(str(path), width=Inches(6.75))
    p = doc.add_paragraph(clean(caption))
    p.style = "Légende"
    p = doc.add_paragraph(clean(observation))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p = doc.add_paragraph(clean(correction))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if reserve:
        p = doc.add_paragraph(clean(reserve))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_phase1_execution_captures(doc: Document, figures: dict[str, Path]) -> None:
    add_capture_analysis_page(
        doc,
        "Dépendance initiale au modèle de langage",
        figures["capture_main"],
        "Figure 4. Capture originale d’une exécution réelle de la branche origin/main.",
        "Le routage vers Gemini intervient pour presque chaque métrique. Lorsque le service renvoie une erreur 503, cout_construction, investissement_total et opex_1 restent à explorer. Le système ne dispose alors d’aucun repli suffisamment robuste pour proposer un candidat exploitable.",
        "La version récente commence par une présélection déterministe, combine TF IDF et embeddings locaux, fusionne les rangs puis applique les preuves structurelles et métier. La liste déterministe est conservée lorsque Gemini échoue. Le LLM n’est sollicité que lorsque l’écart entre candidats demeure insuffisant ou que les preuves sont contradictoires.",
        "Cette capture documente une exécution réelle dans un contexte de forte demande de Gemini. Elle démontre la fragilité de l’architecture initiale, mais ne constitue pas à elle seule une mesure contrôlée de la disponibilité habituelle du service.",
    )
    add_capture_analysis_page(
        doc,
        "Détection insuffisante de la structure des métriques",
        figures["capture_structures"],
        "Figure 5. Capture originale montrant les structures inconnues pour investissement_total et opex_1.",
        "La branche hybride publiée localise des candidats et conserve une shortlist, mais elle ne sait pas encore interpréter la forme de certaines valeurs. Une valeur proposée sans structure fiable ne permet pas de savoir s’il s’agit d’un scalaire, d’un total, d’une décomposition ou d’une série temporelle.",
        "Le code récent sépare l’observation de la classification. Il examine les formats Excel, les en têtes temporels, les mots signalant un total et les formules d’agrégation. L’analyse des dépendances peut requalifier un scalaire en scalar_aggregate lorsqu’une formule prouve qu’il agrège plusieurs cellules. Un résolveur year1_value traite en outre l’OPEX de première année comme une valeur directe lorsque ce périmètre est explicite.",
        "La mention structure inconnue peut encore apparaître dans un classeur atypique. L’amélioration réduit l’incertitude et l’explique, mais une validation sur plusieurs modèles reste nécessaire avant de parler de couverture générale.",
    )
    add_capture_analysis_page(
        doc,
        "Unités absentes et confiance excessive dans des valeurs aberrantes",
        figures["capture_units"],
        "Figure 6. Capture originale montrant des valeurs proposées sans unité et plusieurs résultats économiquement aberrants.",
        "La capture montre qu’un score élevé ne garantissait pas la justesse. Le TRI des fonds propres et le WACC sont proposés à zéro, le productible correspond à une période courte, et les durées de construction et de concession sont rattachées à des cellules dont le sens métier est inadéquat. Les unités ne sont pas présentées, ce qui empêche aussi de vérifier l’ordre de grandeur.",
        "Le collecteur récent recherche les unités dans les lignes supérieures et conserve le contexte autour de la cellule. Le filtre compare ensuite la famille détectée à celle attendue. Les valeurs nulles reçoivent un facteur métier nul. Des règles privilégient le productible annuel, exigent un libellé explicite pour les durées et rejettent les périodes de grâce ou les échéances de dette comme durée de concession. L’analyste peut enfin corriger indépendamment la valeur et l’unité.",
        "La détection d’une unité ne suffit pas encore à normaliser complètement la devise, l’échelle, l’année monétaire et le périmètre économique. Ce chantier reste nécessaire avant toute comparaison financière entièrement automatisée.",
    )
    add_capture_analysis_page(
        doc,
        "Indisponibilité interprétée comme disponibilité",
        figures["capture_availability"],
        "Figure 7. Capture originale montrant une série d’indisponibilité proposée directement comme disponibilité.",
        "Le système avait trouvé la bonne zone du classeur, mais il en restituait le contenu sans comprendre la relation économique. Une indisponibilité de dix pour cent devait conduire à une disponibilité de quatre vingt dix pour cent. L’appel au LLM n’a ni détecté ni corrigé cette inversion.",
        "Un résolveur déterministe spécialisé recherche maintenant les termes unavailability et indisponibilité. Il applique availability = 1 minus unavailability à une valeur simple, à chaque palier d’une série et aux bornes d’une tendance. La valeur source, les segments d’origine et la transformation sont conservés dans le détail du résultat. Un test automatisé vérifie notamment que dix pour cent devient quatre vingt dix pour cent.",
        "La transformation n’est appliquée que lorsque le libellé source établit explicitement qu’il s’agit d’une indisponibilité. Cette condition évite d’inverser arbitrairement une série dont le sens n’est pas démontré.",
    )


def add_equation(doc: Document, equation: str, explanation: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(equation)
    r.font.name = "Cambria Math"
    r.font.size = Pt(12)
    r.italic = True
    q = doc.add_paragraph(clean(explanation))
    q.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_retrieval_theory(doc: Document, figures: dict[str, Path]) -> None:
    title(doc, "Fondements de la recherche lexicale par TF IDF")
    paragraphs = [
        "La localisation d’une métrique dans un modèle financier peut d’abord être abordée comme un problème de recherche d’information. Le système dispose d’une description du concept attendu, par exemple le coût de construction, et doit classer plusieurs milliers de libellés extraits du classeur. TF IDF fournit une première réponse déterministe à ce problème. La méthode accorde un poids élevé aux termes fréquents dans un libellé donné, mais rares dans l’ensemble du catalogue. Un mot très général comme total apporte ainsi moins d’information qu’un terme spécialisé comme EPC, concession ou WACC.",
        "La composante TF, pour term frequency, représente l’importance d’un terme t dans un document d. Dans le contexte du projet, un document correspond au texte enrichi d’un candidat : son libellé, sa section et son contexte proche. La composante IDF, pour inverse document frequency, réduit le poids des mots présents presque partout. Les formes lissées ci dessous évitent une division par zéro et restent stables lorsque le catalogue évolue légèrement.",
    ]
    for text_value in paragraphs:
        p = doc.add_paragraph(clean(text_value)); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_equation(doc, "TF(t,d) = f(t,d) / Σₜ′ f(t′,d)", "f(t,d) désigne le nombre d’occurrences du terme t dans le texte du candidat d. La somme au dénominateur correspond au nombre total de termes retenus dans ce texte.")
    add_equation(doc, "IDF(t) = ln((N + 1) / (df(t) + 1)) + 1", "N représente le nombre de candidats du catalogue et df(t) le nombre de candidats contenant le terme t. Plus un terme est rare dans le classeur, plus sa contribution est importante.")
    add_equation(doc, "TF IDF(t,d) = TF(t,d) × IDF(t)", "Le candidat et la requête sont représentés par des vecteurs de poids TF IDF. Leur proximité est ensuite mesurée à l’aide du cosinus.")
    add_equation(doc, "sim(q,d) = (v(q) · v(d)) / (‖v(q)‖₂ ‖v(d)‖₂)", "Une similarité proche de un indique des distributions lexicales proches. Cette mesure est rapide, reproductible et particulièrement efficace lorsque les termes métiers ou leurs variantes figurent explicitement dans le classeur.")
    p = doc.add_paragraph(clean("TF IDF présente néanmoins une limite importante. Deux expressions peuvent désigner la même réalité sans partager suffisamment de mots. Construction cost, coût des travaux, EPC amount et engineering procurement and construction peuvent être proches sur le plan métier tout en restant éloignés lexicalement. Une recherche exclusivement fondée sur TF IDF risque alors de classer trop bas une cellule pourtant pertinente.")); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    title(doc, "Apport des embeddings sémantiques locaux")
    paragraphs = [
        "Les embeddings complètent la recherche lexicale en représentant un texte par un vecteur dense appris à partir de régularités linguistiques. Une fonction d’encodage transforme une expression en un point d’un espace de dimension m. Des expressions rencontrées dans des contextes comparables ont tendance à être rapprochées dans cet espace, même lorsqu’elles utilisent des mots différents.",
        "Dans le prototype récent, l’encodage est exécuté localement. Cette décision réduit la quantité de données envoyée à un service externe et évite de dépendre d’un appel LLM pour comparer chaque libellé. Les vecteurs du catalogue peuvent en outre être réutilisés, alors qu’un appel distant doit être attendu et facturé à chaque requête.",
    ]
    for text_value in paragraphs:
        p = doc.add_paragraph(clean(text_value)); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_equation(doc, "e(x) = Encodeur(x) ∈ ℝᵐ", "x est le texte associé à une métrique ou à un candidat. L’encodeur produit un vecteur dense de dimension m.")
    add_equation(doc, "simₑ(q,d) = (e(q) · e(d)) / (‖e(q)‖₂ ‖e(d)‖₂)", "La similarité cosinus est également utilisée dans l’espace des embeddings. Elle permet de récupérer des paraphrases et des variantes terminologiques que la recherche lexicale peut manquer.")
    p = doc.add_paragraph(clean("Les embeddings ne remplacent toutefois pas TF IDF. Une proximité sémantique générale peut rapprocher des expressions financièrement différentes, par exemple coût de construction, coût total du projet et coût de financement. Les acronymes rares, les références précises à une tranche de dette et les conventions propres à un classeur sont parfois mieux traités par la correspondance lexicale. La représentation dense apporte donc du rappel, mais peut introduire des voisins sémantiques trop larges.")); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    title(doc, "Pourquoi combiner TF IDF et embeddings")
    paragraphs = [
        "Le choix d’une architecture combinée vient de la complémentarité de leurs erreurs. TF IDF est précis lorsque le vocabulaire du référentiel apparaît dans le classeur. Les embeddings sont utiles lorsque le promoteur emploie une paraphrase, une autre langue ou une nomenclature différente. Conserver les deux listes évite de demander à une seule méthode de résoudre simultanément la correspondance exacte et la proximité de sens.",
        "Le système récent produit donc plusieurs classements indépendants : une présélection fondée sur les synonymes métier, un classement TF IDF et un classement par embeddings. Ces listes sont réunies par une fusion réciproque des rangs. Cette méthode se fonde sur la position d’un candidat plutôt que sur la valeur brute de scores qui ne sont pas directement comparables. Un candidat bien classé par plusieurs méthodes reçoit naturellement une priorité supérieure.",
    ]
    for text_value in paragraphs:
        p = doc.add_paragraph(clean(text_value)); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_equation(doc, "RRF(d) = Σᵣ 1 / (k + rangᵣ(d))", "r parcourt les méthodes de recherche, rangᵣ(d) est la position du candidat d dans la liste r et k est une constante d’amortissement. Un candidat absent d’une liste n’y reçoit aucune contribution.")
    p = doc.add_paragraph(clean("Après la fusion, le classement n’est pas encore considéré comme une décision finale. Le pipeline examine la structure de la cellule, son unité, sa plage plausible, sa formule et son adéquation au sens métier. Pour les métriques qui doivent représenter un total, l’analyse des dépendances peut ajouter un signal lorsqu’une formule agrège plusieurs cellules. Les règles métier peuvent au contraire ramener le score à zéro lorsqu’une valeur est nulle ou lorsque le libellé désigne clairement un autre périmètre.")); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_equation(doc, "Sᵣₑₜ(d) = normaliser(RRF(d)) + B_dépendances(d)", "Le score de récupération rassemble la fusion des rangs et, lorsque le concept attend un agrégat, un bonus fondé sur les dépendances de formule.")
    add_equation(doc, "Sfinal(d) = Sstructure(d) × Fmétier(d)", "Le diagnostic structurel est pondéré par un facteur métier compris entre zéro et un. Une incompatibilité déterministe forte peut ainsi éliminer un candidat avant tout appel coûteux au LLM.")
    title(doc, "Place résiduelle du modèle de langage")
    paragraphs = [
        "La combinaison de TF IDF, des embeddings et des règles structurelles ne vise pas à supprimer toute interprétation sémantique. Elle vise à réserver le modèle de langage aux cas pour lesquels son apport est réellement utile. Lorsque le premier candidat se détache nettement et que la valeur, l’unité, la structure et le périmètre convergent, le pipeline peut conserver la décision déterministe. Lorsque plusieurs cellules restent proches ou que les preuves se contredisent, le modèle reçoit une liste courte et fermée de candidats.",
        "Cette organisation répond directement aux difficultés visibles dans les captures. Dans la branche initiale, une erreur 503 interrompait pratiquement la recherche. Dans la branche hybride publiée, une liste déterministe survivait à l’erreur, mais le LLM était encore appelé pour chaque métrique. L’architecture actuelle ajoute un routage conditionnel afin de réduire ces appels. Le LLM ne peut pas inventer une cellule : il doit choisir dans la liste transmise et sa réponse est vérifiée par le code.",
        "Le choix combiné possède enfin un intérêt expérimental. Il permet de mesurer séparément la contribution de la recherche lexicale, de la recherche sémantique, des règles structurelles et du reranking éventuel par LLM. Une étude d’ablation devra comparer TF IDF seul, embeddings seuls, fusion des deux, puis fusion enrichie par les règles métier. Les résultats devront porter sur plusieurs classeurs avant de conclure que le gain observé se généralise au delà du modèle actuellement disponible.",
    ]
    for text_value in paragraphs:
        p = doc.add_paragraph(clean(text_value)); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
def add_architecture_chapter(doc: Document, figure_path: Path) -> None:
    title(doc, "Architecture globale détaillée du système")
    p = doc.add_paragraph(clean(
        "L’architecture actuelle est le résultat direct des difficultés rencontrées pendant les premières expérimentations. Le système initial concentrait une part trop importante de l’interprétation dans les appels au modèle de langage. Cette organisation paraissait simple, mais elle rendait l’extraction lente, sensible aux erreurs de service et difficile à reproduire. La version actuelle répartit explicitement les responsabilités entre trois ensembles : la Phase 1 extrait et fait valider les métriques du classeur, la banque de benchmarks entretient les références externes, et la Phase 2 réalise les calculs, les comparaisons et la production du rapport. Des mécanismes transversaux assurent la provenance, les tests, l’instrumentation et l’enregistrement des décisions humaines."
    )); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_picture(str(figure_path), width=Inches(6.8))
    p = doc.add_paragraph("Figure 2. Architecture cible détaillée distinguant les composants déjà conçus des fonctions prévues."); p.style = "Légende"
    for text_value in [
        "La lecture de cette figure repose sur une distinction volontaire entre l’existant et la cible. Les blocs à contour continu représentent les mécanismes déjà conçus ou intégrés au prototype. Les blocs à contour discontinu correspondent aux fonctions prévues, dont le rôle est défini mais dont l’intégration complète reste à réaliser et à évaluer. La couleur orangée identifie les interventions du modèle de langage. Cette convention évite de présenter comme acquis un développement qui relève encore de la prochaine étape du projet.",
        "L’architecture cible ne réduit pas le modèle de langage à un simple mécanisme transversal. Elle lui attribue trois interventions fonctionnelles distinctes, placées à des moments où une compréhension linguistique apporte une valeur réelle, tandis que le code conserve les calculs, les règles de contrôle et la traçabilité.",
    ]:
        p = doc.add_paragraph(clean(text_value)); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title(doc, "Séparation des responsabilités")
    for text_value in [
        "Le principe structurant est que le code conserve la conduite du traitement. Il décide quelles étapes doivent être exécutées, applique les règles de filtrage, effectue les transformations numériques et produit les artefacts. Le modèle de langage n’ordonne pas librement le pipeline et ne réalise pas les calculs financiers. Selon l’étape concernée, il produit une explication terminologique, départage une ambiguïté bornée ou enrichit l’interprétation de faits déjà calculés. Cette séparation réduit le risque qu’une variation de réponse modifie silencieusement l’ordre des opérations ou la définition d’un indicateur.",
        "L’analyste constitue le troisième acteur de l’architecture. Il valide une cellule proposée, choisit une alternative, corrige une valeur, modifie une unité ou déclare qu’une métrique est indisponible. Ces actions ne sont pas incorporées au fichier Excel d’origine. Elles sont enregistrées dans un registre distinct avec leur provenance. Le système peut ainsi distinguer ce qui vient du promoteur, ce qui a été interprété automatiquement et ce qui résulte d’une décision humaine.",
        "Cette organisation est particulièrement importante dans un contexte réglementaire. Une conclusion ne doit pas seulement être plausible : elle doit pouvoir être reproduite et expliquée. Le registre des hypothèses, les sélections de projets comparables et le manifeste d’analyse forment donc des frontières explicites entre les étapes. Une phase ultérieure peut être rejouée sans relancer toutes les opérations antérieures lorsque ses entrées n’ont pas changé.",
    ]:
        p = doc.add_paragraph(clean(text_value)); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title(doc, "Flux de la Phase 1")
    for text_value in [
        "La Phase 1 reçoit un classeur Excel non standardisé et le référentiel ARSEL. Le classeur est ouvert sans modification. Le collecteur parcourt séquentiellement les feuilles retenues et rassemble les libellés, les premières valeurs numériques voisines, la section, le contexte textuel et les éventuelles unités situées au dessus des cellules. La lecture séquentielle évite les nombreux accès dispersés qui avaient fortement ralenti certaines versions antérieures.",
        "Pour chaque concept attendu, plusieurs moteurs de recherche produisent ensuite leurs candidats. La présélection utilise les synonymes et les expressions définis dans le référentiel. TF IDF privilégie les correspondances lexicales informatives. Les embeddings locaux récupèrent les expressions sémantiquement proches. La fusion réciproque des rangs rassemble ces listes sans supposer que leurs scores bruts partagent la même échelle.",
        "La récupération est suivie d’une analyse plus exigeante. Le système caractérise la cellule et sa formule, recherche une structure scalaire, agrégée ou temporelle, compare la famille d’unité à celle attendue et vérifie la plage plausible. Les règles métier rejettent les périmètres manifestement incompatibles. Lorsqu’un concept attend un total, l’analyse des dépendances examine si la cellule résulte effectivement d’une agrégation. Cette seconde couche évite qu’une simple proximité de vocabulaire suffise à établir une proposition.",
        "Le routage vers le LLM intervient seulement lorsque les preuves déterministes ne permettent pas de départager suffisamment les candidats. Le modèle reçoit une liste courte et ne peut choisir qu’une cellule qui y figure. Une réponse qui invente une adresse, omet les champs attendus ou arrive après une indisponibilité du service est rejetée. La shortlist déterministe reste alors disponible pour l’analyste. À l’issue de la validation, le fichier hypotheses_validees.json conserve la valeur, l’unité, la source et la décision associée à chaque métrique.",
    ]:
        p = doc.add_paragraph(clean(text_value)); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title(doc, "Architecture de la banque de benchmarks")
    for text_value in [
        "La banque de benchmarks est indépendante d’un dossier particulier. Elle reçoit actuellement deux catégories de données. La Banque mondiale fournit des projets individuels issus de la base PPI. L’IRENA fournit des statistiques sectorielles structurées par technologie, année, géographie et métrique. Cette séparation évite d’interpréter une moyenne sectorielle comme s’il s’agissait d’un projet comparable, ou inversement.",
        "L’orchestrateur périodique surveille les sources officielles, détecte les nouvelles versions et télécharge d’abord les fichiers dans un emplacement temporaire. Le contenu est identifié par un checksum SHA 256. Une version brute déjà enregistrée n’est pas écrasée. Les adaptateurs vérifient le format, reconnaissent les feuilles attendues et transforment les données vers les schémas canoniques des sources, projets, observations et événements de normalisation.",
        "L’ingestion alimente d’abord une base de staging. Les contrôles vérifient notamment la présence d’une technologie, la cohérence des unités, les doublons et la conformité des objets. La base active n’est jamais remplacée automatiquement à la fin d’un téléchargement. Une promotion explicite est nécessaire après examen du rapport qualité. Cette décision protège les analyses contre une modification imprévue du format d’une source ou contre une ingestion techniquement valide mais économiquement incohérente.",
        "Dans DuckDB, les projets individuels et les observations sectorielles conservent leurs identifiants, leur source, leur période, leur technologie, leur géographie, leur unité brute et leur représentation normalisée. Les événements de normalisation permettent de comprendre comment une valeur a été transformée. Les snapshots et les checksums permettent de retrouver le fichier exact à l’origine d’une observation utilisée dans un rapport.",
    ]:
        p = doc.add_paragraph(clean(text_value)); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title(doc, "Flux de la Phase 2")
    for text_value in [
        "La Phase 2 ne retourne pas chercher librement des valeurs dans le classeur. Elle consomme le registre validé et le contexte du projet. Sa première responsabilité est de normaliser les représentations : nombres formatés, pourcentages, listes et séries temporelles. Elle calcule ensuite les indicateurs dérivés selon des formules déterministes. Cette règle garantit qu’un même registre produit les mêmes ratios, indépendamment d’une variation du modèle de langage.",
        "Le profil du projet est utilisé pour rechercher des projets comparables. La technologie constitue un filtre strict. Les candidats restants sont évalués selon la géographie, la capacité, la période, la durée contractuelle, la structure de financement et la disponibilité des métriques. L’analyste approuve ou rejette les projets proposés. Les statistiques de petite taille, comme le minimum, les quartiles et la médiane, ne sont calculées que sur les projets approuvés et séparément pour chaque métrique disponible.",
        "Les statistiques IRENA suivent une autre voie. Elles apportent une référence sectorielle pour une technologie et une géographie données. Le moteur vérifie l’unité et le périmètre avant de positionner la valeur du projet. Une référence peut être conservée comme contexte sans produire de verdict lorsque la comparabilité est insuffisante. Cette prudence évite de comparer directement un investissement incluant des frais financiers à un coût installé qui les exclut.",
        "Les constats déterministes, les écarts, les risques et les données manquantes alimentent ensuite l’analyse professionnelle. Les sorties JSON et Markdown conservent la structure détaillée, tandis que le document Word fournit une version destinée à la lecture. Le manifeste final enregistre les chemins, les tailles et les checksums des principales entrées et sorties. Il devient ainsi possible de démontrer quelle base de benchmarks, quel registre et quel contexte ont servi à produire un rapport donné.",
    ]:
        p = doc.add_paragraph(clean(text_value)); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title(doc, "Mécanismes transversaux et points de contrôle")
    for text_value in [
        "La provenance traverse toutes les couches. Dans la Phase 1, elle relie la proposition à une feuille, une cellule, une formule et une unité détectée. Dans la banque, elle relie une observation à une source et à un snapshot. Dans la Phase 2, elle relie un indicateur dérivé aux métriques validées qui entrent dans son calcul. Cette continuité est indispensable pour répondre à une contestation sans reconstruire manuellement le raisonnement.",
        "L’instrumentation des appels LLM enregistre leur nombre, leur latence et leur statut. Les traces historiques ont montré que la définition du succès devait elle même être précisée : un pipeline peut retourner une shortlist malgré l’échec du fournisseur. La version récente distingue donc conceptuellement la réussite du service externe, la validité de sa réponse et la capacité du pipeline à poursuivre par un repli déterministe.",
        "Les tests automatisés couvrent les schémas, les règles métier, la navigation de l’analyste, les transformations particulières et la banque de benchmarks. Ils ne remplacent pas l’évaluation sur des classeurs réels, mais ils protègent les comportements déjà explicités. L’inversion de l’indisponibilité, le rejet d’une valeur nulle et la préférence pour un productible annuel constituent des exemples de connaissances métier transformées en règles vérifiables.",
        "Enfin, l’architecture reste volontairement évolutive. Un modèle local peut remplacer le fournisseur infonuagique derrière l’interface prévue, une nouvelle source peut être ajoutée par un adaptateur, et une interface graphique pourra consommer les mêmes objets de validation. Ces évolutions ne doivent pas modifier les formules financières ni les frontières de provenance. Le noyau déterministe reste la source des valeurs et des calculs, mais le modèle de langage possède bien des responsabilités fonctionnelles identifiées dans l’architecture cible.",
    ]:
        p = doc.add_paragraph(clean(text_value)); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title(doc, "Les trois interventions prévues du modèle de langage")
    for text_value in [
        "La première intervention se situe à l’étape zéro, avant la validation des métriques. Le référentiel ARSEL peut déjà fournir de manière déterministe la définition des métriques connues, puisque leur sens doit rester stable d’un dossier à l’autre. Cette base ne suffit cependant pas à couvrir toutes les expressions propres à un modèle financier. Dans le classeur Kikot, des termes comme EPC ou WHT peuvent être familiers à un analyste expérimenté, mais ils ne le seront pas nécessairement pour un utilisateur plus récent. À terme, le système devra repérer les termes techniques absents du référentiel, puis demander au modèle de langage de produire une définition contextualisée, concise et adaptée au projet. Cette fonction est représentée en pointillé parce qu’elle est prévue mais n’est pas encore entièrement intégrée.",
        "La deuxième intervention concerne la sélection des métriques en Phase 1. Elle existe déjà sous une forme bornée. TF IDF, les embeddings, la fusion des rangs et les preuves structurelles construisent d’abord une liste fermée de cellules candidates. Le modèle de langage n’est sollicité que lorsqu’une ambiguïté réelle subsiste et il ne peut pas proposer une cellule extérieure à cette liste. Son rôle est donc celui d’un arbitre sémantique contrôlé, et non celui d’un moteur de recherche autonome dans le classeur.",
        "La troisième intervention est prévue après les calculs déterministes de la Phase 2. Le système devra d’abord normaliser les valeurs, calculer les indicateurs, sélectionner les comparables et quantifier les écarts par rapport aux projets et aux statistiques sectorielles. Il constituera ensuite un ensemble de faits structuré comprenant les résultats, les sources, les risques, les limites et les données manquantes. Le modèle de langage pourra s’appuyer sur cet ensemble pour ajouter du contexte, expliquer les interactions entre plusieurs indicateurs et approfondir l’interprétation financière. Il ne recalculera pas les valeurs et ne pourra pas remplacer les chiffres déterministes par ses propres estimations.",
        "Cette répartition répond à un objectif simple : employer le modèle de langage là où la compréhension du vocabulaire et la mise en relation des faits sont utiles, sans lui confier les opérations qui exigent une stricte reproductibilité. Les définitions générées et l’analyse enrichie devront rester contrôlables par l’analyste et clairement distinguées des résultats calculés par le système.",
    ]:
        p = doc.add_paragraph(clean(text_value)); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title(doc, "Parcours complet d’un dossier dans l’architecture")
    for text_value in [
        "Le traitement d’un dossier commence par l’identification du classeur reçu et de son contexte. Le fichier n’est pas converti vers un gabarit imposé et aucune formule n’est réécrite. Le système calcule éventuellement les informations nécessaires à la traçabilité du fichier, ouvre une vue contenant les valeurs évaluées et une autre contenant les formules, puis vérifie quelles feuilles attendues sont réellement présentes. Cette double lecture est nécessaire : la valeur calculée sert à l’analyse financière, tandis que la formule fournit des indices sur le rôle structurel de la cellule.",
        "Le collecteur parcourt ensuite les feuilles de manière séquentielle. Lorsqu’il trouve un libellé admissible, il recherche la première valeur numérique située à droite sur la même ligne et conserve les textes environnants. Le contexte vertical est tout aussi important que la ligne elle même, car les modèles financiers placent souvent la devise ou l’échelle dans un en tête commun à plusieurs lignes. Le résultat n’est pas encore une métrique validée, mais un catalogue de couples libellé valeur enrichis de leur adresse, de leur section et de leur contexte.",
        "Pour le coût de construction, par exemple, le référentiel fournit des synonymes, une définition, une nature monétaire, une structure agrégée attendue et des signaux de périmètre. Les moteurs de récupération interrogent le catalogue avec ces informations. Une cellule contenant EPC peut être favorisée lexicalement, tandis qu’une cellule intitulée engineering procurement and construction peut être récupérée par proximité sémantique. Une ligne de frais financiers peut partager le mot total, mais les règles de périmètre et d’unité doivent l’empêcher de dominer le classement.",
        "Les premiers candidats sont ensuite caractérisés à partir du classeur. Le détecteur examine le format numérique, la présence d’une formule, l’existence d’un en tête temporel et les mots signalant un total ou un sélecteur. L’index de dépendances suit les références de formule dans une profondeur bornée afin d’estimer si la cellule rassemble plusieurs feuilles terminales. Cette analyse ne cherche pas à recalculer Excel. Elle fournit des preuves structurelles qui complètent la similarité textuelle.",
        "Lorsque le premier candidat présente une avance suffisante et qu’aucune incompatibilité forte n’est détectée, la proposition peut être préparée sans solliciter Gemini. Si deux cellules restent proches, le modèle reçoit les informations utiles concernant une liste courte. Son rôle consiste à départager des options existantes, non à parcourir le classeur ou à fabriquer une adresse. Quelle que soit l’issue, l’analyste voit la proposition, la structure, l’unité, les signaux favorables et les alternatives. Il peut revenir à la métrique précédente si une décision doit être corrigée.",
        "Après la validation des vingt deux métriques, le registre devient l’entrée contractuelle de la Phase 2. Les corrections apportées par l’analyste sont distinguées des valeurs extraites. La Phase 2 normalise les représentations puis construit le profil du projet. Pour un projet hydroélectrique, ce profil peut contenir la puissance, le pays, l’année, le coût d’investissement, la durée de concession et la structure de financement. Une donnée absente reste explicitement absente et ne doit pas être inventée pour compléter le profil.",
        "Le profil interroge ensuite la base active de benchmarks. Les projets individuels incompatibles sur la technologie sont éliminés. Les projets restants reçoivent un classement explicable et sont présentés à l’analyste. L’approbation de chaque comparable est conservée dans un fichier dont le checksum dépend du registre, du contexte et de la base utilisée. Si l’une de ces entrées change, la sélection antérieure est signalée comme obsolète plutôt que réutilisée silencieusement.",
        "Les calculs financiers, les contrôles internes, les statistiques de projets pairs et les positions sectorielles sont enfin réunis dans une analyse structurée. Le générateur Word n’invente pas de nouveaux résultats : il organise les faits déjà produits. Le manifeste clôt le traitement en enregistrant les artefacts. Ce parcours complet permet de revenir depuis une conclusion jusqu’au benchmark, au calcul, à la métrique validée et à la cellule d’origine lorsque toutes les informations nécessaires sont disponibles.",
    ]:
        p = doc.add_paragraph(clean(text_value)); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title(doc, "Modes dégradés et comportement en cas d’échec")
    for text_value in [
        "L’architecture ne considère pas l’indisponibilité d’un composant externe comme une situation exceptionnelle impossible à gérer. Si le modèle de langage est indisponible, la Phase 1 conserve les candidats déterministes et transmet l’ambiguïté à l’analyste. Si les embeddings locaux ne peuvent pas être chargés, le système revient à TF IDF et signale ce repli. La qualité potentielle peut diminuer, mais le traitement ne doit pas se transformer en résultat faussement certain.",
        "Une autre catégorie d’échec concerne les données du classeur. Une unité peut être absente, une formule peut référencer une zone non prise en charge ou une structure temporelle peut être trop irrégulière. Dans ces cas, le système doit conserver la valeur brute, le diagnostic incomplet et les raisons de l’incertitude. La correction par l’analyste demeure possible, mais son origine doit être enregistrée. Cette stratégie est préférable à une conversion automatique fondée sur une hypothèse non démontrée.",
        "La banque de benchmarks suit le même principe. Si l’IRENA modifie le nom du fichier ou l’organisation des feuilles, l’adaptateur tente une reconnaissance dynamique. Lorsque le format ne correspond à aucun schéma connu, l’ingestion échoue proprement et la base active reste intacte. Une nouvelle source n’est donc jamais promue simplement parce que son téléchargement a réussi. Le rapport de contrôle doit expliquer le point de rupture et permettre une adaptation revue du connecteur.",
        "La Phase 2 refuse également certaines comparaisons. Deux valeurs peuvent partager une apparence numérique tout en étant économiquement différentes. Un coût en milliers d’euros ne doit pas être comparé à un montant en euros sans facteur d’échelle. Un coût de construction hors financement ne doit pas recevoir le même périmètre qu’un investissement total incluant les intérêts intercalaires. Lorsque les métadonnées nécessaires manquent, le système doit produire un statut non comparable ou comparable avec prudence, et non une conclusion artificiellement précise.",
        "Ces modes dégradés répondent à une exigence centrale : la continuité du traitement ne doit jamais être obtenue au prix d’une perte silencieuse de sens. Une proposition incomplète mais explicitement qualifiée est plus utile à l’analyste qu’une valeur présentée avec une confiance élevée sans justification. Les captures historiques montrent précisément pourquoi cette distinction a été introduite : la présence d’un score de cent pour cent n’empêchait pas certaines valeurs nulles ou certains mauvais périmètres d’être proposés.",
    ]:
        p = doc.add_paragraph(clean(text_value)); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title(doc, "Contrats de données entre les composants")
    for text_value in [
        "La séparation en sous systèmes n’est utile que si les échanges sont définis. Entre la Phase 1 et la Phase 2, le contrat principal est le registre des hypothèses. Chaque entrée doit identifier la métrique, sa valeur, son unité, sa source et son statut. Les structures complexes devront progressivement préciser si la donnée est un scalaire, une liste, une série temporelle ou un ensemble de tranches. Cette information empêche la Phase 2 de réduire arbitrairement une trajectoire à son premier nombre.",
        "Entre la banque de benchmarks et la Phase 2, les observations doivent conserver la valeur brute et la valeur normalisée, l’unité, la devise, l’année monétaire, la technologie, la géographie et le périmètre. Le stockage d’un chiffre seul serait insuffisant. La comparabilité repose en effet autant sur ces métadonnées que sur la valeur elle même. Les schémas canoniques constituent ainsi une protection contre les rapprochements opportunistes entre sources hétérogènes.",
        "Entre les calculs et la couche narrative future, le contrat devra être encore plus strict. Le modèle de langage recevra des faits structurés, des écarts calculés, des benchmarks approuvés et des données manquantes. Il pourra enrichir l’explication ou proposer des causes à examiner, mais il ne devra pas modifier les résultats numériques. Chaque affirmation narrative importante devra être rattachée à un identifiant de fait ou être clairement présentée comme une hypothèse analytique.",
        "Enfin, l’interface utilisateur prévue ne devra pas contourner ces contrats. Une vue de type tableau de bord pourra afficher simultanément les candidats et faciliter la comparaison, mais les actions continueront à produire les mêmes instructions validées que la console. Cette continuité permettra de changer l’ergonomie sans réécrire le moteur et sans créer deux chemins de décision incompatibles.",
    ]:
        p = doc.add_paragraph(clean(text_value)); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title(doc, "Justification des principaux choix techniques")
    for text_value in [
        "Python a été retenu comme langage d’intégration parce qu’il réunit les bibliothèques nécessaires à la lecture d’Excel, à la recherche d’information, au calcul financier, à la manipulation de données et à la génération de documents. Ce choix réduit le nombre de technologies à maintenir et facilite l’écriture de tests. Il ne signifie pas que toutes les opérations doivent être réalisées dans un seul module. La restructuration récente sépare au contraire les responsabilités entre arsel_core, benchmark_bank et phase2 afin de limiter les dépendances circulaires et de permettre une évolution indépendante des composants.",
        "OpenPyXL est utilisé pour accéder aux valeurs, aux formats et aux formules du classeur sans imposer l’installation d’Excel dans le chemin principal d’extraction. Deux modes de lecture sont complémentaires. Le mode data_only fournit les dernières valeurs calculées enregistrées dans le fichier, tandis que la vue des formules permet de détecter les agrégations et les dépendances. Cette approche comporte une limite connue : OpenPyXL ne recalcule pas le classeur. Une valeur mise en cache peut être ancienne si le promoteur n’a pas enregistré le fichier après recalcul. Ce risque devra être signalé ou contrôlé dans une version opérationnelle.",
        "La recherche TF IDF repose sur une représentation creuse et peu coûteuse. Les embeddings produisent une représentation dense plus expressive, mais leur chargement et leur calcul peuvent être plus lourds. L’architecture conserve donc un mécanisme de repli : l’échec de l’index sémantique ne bloque pas la recherche lexicale. La mise en cache des représentations locales devra être mesurée avec soin, car une comparaison des temps n’est interprétable que si l’on distingue le premier chargement du modèle des requêtes suivantes.",
        "DuckDB a été choisi pour la banque de benchmarks parce qu’il fournit un moteur analytique embarqué, interrogeable localement et adapté aux volumes actuels du projet. La base peut être distribuée comme un fichier tout en conservant des requêtes SQL, des vues et des transactions. Ce choix simplifie le prototype et limite les transferts de données. Si plusieurs analystes doivent ultérieurement écrire simultanément dans la banque, une architecture de service ou une base serveur pourra devenir nécessaire. Les schémas canoniques devraient néanmoins rester réutilisables.",
        "Les formats JSON jouent le rôle de contrats lisibles entre les étapes. Ils facilitent le contrôle, la comparaison de versions et la reprise d’une analyse. Markdown fournit une sortie intermédiaire facile à examiner, tandis que Word répond au besoin de diffusion institutionnelle. La multiplication de ces formats n’est utile que parce qu’ils sont produits à partir du même objet d’analyse. Trois générateurs indépendants risqueraient autrement de faire apparaître des conclusions divergentes.",
        "Enfin, les checksums SHA 256 ne constituent pas une mesure de qualité des données. Ils garantissent seulement l’identité du contenu. Leur valeur architecturale vient de leur association aux snapshots, aux sélections et au manifeste. Ils permettent de démontrer que deux analyses ont utilisé le même registre ou la même base, puis de concentrer l’examen des différences sur le code, la configuration ou les décisions de l’analyste.",
    ]:
        p = doc.add_paragraph(clean(text_value)); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title(doc, "Évaluation expérimentale attendue de l’architecture")
    for text_value in [
        "L’architecture ne pourra être considérée comme validée sur la seule base du modèle Kikot. L’évaluation devra réunir plusieurs classeurs présentant des organisations, des langues, des technologies et des conventions d’unités différentes. Pour chaque métrique, une vérité terrain devra préciser la cellule correcte, la valeur attendue, l’unité, le périmètre et les alternatives éventuellement acceptables. Cette granularité est nécessaire, car deux cellules peuvent contenir le même nombre tout en représentant des concepts différents.",
        "La première famille de mesures portera sur la récupération. Le rappel à cinq et à dix indiquera si la bonne cellule apparaît dans la shortlist. La précision en première position mesurera la fréquence à laquelle elle est proposée directement. Le rang réciproque moyen décrira la position générale de la vérité terrain. Ces mesures devront être calculées séparément pour TF IDF, les embeddings, leur fusion et la chaîne complète afin de quantifier la contribution de chaque composant.",
        "La deuxième famille portera sur la fidélité finale. Elle distinguera la cellule, la valeur, l’unité et la structure. Une extraction ne sera pas déclarée entièrement correcte si la cellule est juste mais si l’échelle monétaire est absente. L’évaluation devra également mesurer la calibration de la confiance : parmi les propositions affichées à quatre vingt dix pour cent ou davantage, quelle proportion est réellement correcte ? Les captures historiques montrent pourquoi cette question est aussi importante que le rappel brut.",
        "La troisième famille concernera les performances et la robustesse. Le protocole enregistrera le temps de collecte, le chargement des embeddings, le temps par métrique, le nombre d’appels LLM et les erreurs du fournisseur. Des exécutions répétées permettront de séparer le coût du premier chargement de celui du cache chaud. Des scénarios de panne volontaire vérifieront que le repli TF IDF, la shortlist déterministe et l’arrêt propre des ingestions externes fonctionnent comme prévu.",
        "Enfin, l’utilité pour l’analyste devra être évaluée. Le temps consacré à la validation, le nombre de corrections, la compréhension des justifications et la pertinence des projets comparables fourniront des indicateurs plus proches de l’usage réel. Une proposition très précise mais difficile à expliquer ne répondrait qu’imparfaitement au contexte réglementaire. L’objectif final reste de réduire le travail répétitif tout en améliorant la qualité et la traçabilité de l’instruction.",
    ]:
        p = doc.add_paragraph(clean(text_value)); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def parse_source() -> list[tuple[str, list[str]]]:
    text = SOURCE_MD.read_text(encoding="utf 8")
    sections: list[tuple[str, list[str]]] = []
    current = None
    paras: list[str] = []
    for block in re.split(r"\n\s*\n", text):
        block = block.strip()
        if not block:
            continue
        if block.startswith("## "):
            if current:
                sections.append((current, paras))
            current = block[3:].strip()
            paras = []
        elif not block.startswith("# "):
            paras.append(re.sub(r"\s+", " ", block))
    if current:
        sections.append((current, paras))
    return sections


def add_cover(doc: Document) -> None:
    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RAPPORT D’AVANCEMENT DÉTAILLÉ"); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Système interactif d’analyse assistée des projections financières de projets énergétiques"); r.font.size = Pt(18); r.font.color.rgb = RGBColor.from_string(MID_BLUE)
    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    data = [("Étudiant", "Michel Wilfred Essono"), ("Programme", "Maîtrise en informatique, Université de Montréal"), ("Superviseure", "Pre Esma Aïmeur"), ("Organisation d’accueil", "Agence de Régulation du Secteur de l’Électricité du Cameroun"), ("Date", "31 août 2026")]
    for row, (a, b) in zip(table.rows, data):
        set_cell_text(row.cells[0], a, True, BLUE, 10); set_cell_text(row.cells[1], b, False, None, 10); shade(row.cells[0], LIGHT_BLUE)
    doc.add_paragraph()
    callout(doc, "Portée du document", "Le présent rapport décrit l’évolution réelle du prototype, les choix d’architecture, les résultats déjà obtenus, les limites observées et les travaux restant à mener.")
    doc.add_page_break()

    title(doc, "Mise en contexte et retour sur le plan de cadrage")
    context_paragraphs = [
        "Le projet a été défini dans un premier document de cadrage daté du 26 juin 2026. Il portait sur la conception d’un système interactif combinant un modèle de langage et un moteur de calcul déterministe pour assister l’analyse des projections financières de projets hydroélectriques soumis à l’Agence de Régulation du Secteur de l’Électricité du Cameroun. Le point de départ était un besoin opérationnel concret : l’instruction des modèles financiers mobilise une expertise spécialisée, alors que les classeurs reçus des promoteurs ne suivent pas un format commun et peuvent contenir des milliers de cellules, de formules et d’hypothèses réparties entre de nombreuses feuilles.",
        "Le cadrage initial insistait sur le fait que l’outil ne devait pas fonctionner comme une boîte noire transformant directement un fichier Excel en rapport. Dans le contexte réglementaire d’ARSEL, une conclusion peut être discutée par le promoteur et doit pouvoir être justifiée jusqu’aux données qui l’ont produite. Le système devait donc préserver le classeur original, séparer les calculs des interprétations, enregistrer les interventions de l’analyste et permettre de reprendre une analyse sans perdre la provenance des valeurs. Cette exigence distinguait dès le départ le projet d’un simple exercice de génération automatique de texte.",
        "La problématique formulée dans le plan de cadrage reposait sur trois contraintes liées. La première était l’hétérogénéité des entrées : les noms des feuilles, les libellés et la disposition des informations changent selon les promoteurs. La deuxième était l’exigence de déterminisme et de reproductibilité propre à la régulation. La troisième était la nécessité de conserver une intervention humaine en cours de traitement, puisque certaines informations pertinentes ne figurent pas dans le classeur et que l’analyste demeure responsable de la décision. Le projet devait donc concilier souplesse d’interprétation, stabilité des calculs et traçabilité des corrections.",
        "Trois questions de recherche structuraient le travail. La première portait sur la séparation entre le raisonnement du modèle de langage, les calculs du moteur déterministe et les interventions encadrées de l’analyste. La deuxième concernait la découverte de structure et la localisation fiable de concepts financiers dans des classeurs hétérogènes. La troisième prévoyait une comparaison entre des modèles infonuagiques et des modèles locaux quantifiés, en considérant la fidélité, la reproductibilité, le coût et les contraintes de confidentialité. Ces questions demeurent pertinentes, même si leur traduction technique s’est précisée au cours de la réalisation.",
        "L’architecture envisagée au départ plaçait le modèle de langage dans un rôle important de découverte sémantique, tout en réservant les calculs financiers à Python. Des points d’arrêt devaient permettre à l’analyste de valider les correspondances, d’ajouter une information absente et de demander une explication. Les livrables prévoyaient un prototype interactif, un registre des écarts et des origines, un protocole d’évaluation, une comparaison de modèles, un rapport de stage et une démonstration aux analystes d’ARSEL. Le plan estimait que la découverte de structure et l’extraction constitueraient les premières étapes avant la consolidation du moteur d’analyse et de la couche d’explication.",
        "Les premières expérimentations ont confirmé la pertinence du besoin, mais elles ont également montré qu’un recours trop fréquent au modèle de langage rendait la Phase 1 lente et fragile. Les erreurs de service et les limites de quota pouvaient interrompre la recherche ou laisser la majorité des métriques sans proposition. À l’inverse, une sélection exclusivement lexicale ne couvrait pas suffisamment les variations de vocabulaire. Le projet a donc évolué vers une architecture hybride dans laquelle TF IDF, les embeddings locaux, les preuves structurelles et les règles métier construisent d’abord une liste explicable. Le modèle de langage est progressivement devenu un mécanisme de résolution des ambiguïtés plutôt que le centre du pipeline.",
        "Un autre élargissement majeur concerne les benchmarks. Le cadrage initial mentionnait la comparaison avec des projets soumis à des contraintes proches, mais ne définissait pas encore une infrastructure complète pour constituer et gouverner ces références. La réalisation a conduit à créer une banque indépendante dans DuckDB, alimentée par les projets de la Banque mondiale et les statistiques sectorielles de l’IRENA. Les snapshots immuables, le staging, les contrôles qualité et la promotion humaine répondent à la même exigence de traçabilité que celle appliquée au classeur du promoteur.",
        "L’évolution de l’architecture ne constitue donc pas un abandon du plan de cadrage. Elle correspond à une concrétisation de ses principes. La séparation entre interprétation et calcul est devenue plus stricte, l’intervention de l’analyste est mieux définie, et la provenance s’étend désormais jusqu’aux sources externes. Certaines ambitions restent toutefois à achever, notamment l’évaluation sur plusieurs classeurs, la comparaison entre modèles locaux et infonuagiques, la normalisation financière complète et le dialogue explicatif à plusieurs niveaux de détail. Le présent rapport doit être lu comme le passage d’une intention de recherche à une première chaîne fonctionnelle, puis à une phase de consolidation scientifique et opérationnelle.",
    ]
    for text_value in context_paragraphs:
        p = doc.add_paragraph(clean(text_value))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    title(doc, "Sommaire détaillé")
    items = [
        "1. Résumé exécutif et état général", "2. Rappel du cadrage initial", "3. Architecture globale détaillée",
        "4. Évolution de l’architecture", "5. Phase 1 : extraction hybride des métriques", "6. Fondements de TF IDF et des embeddings",
        "7. Validation par l’analyste et traçabilité", "8. Banque de benchmarks et gouvernance DuckDB",
        "9. Sélection des projets comparables et statistiques", "10. Phase 2 : calculs, comparaisons et rapport",
        "11. Positionnement par rapport aux questions de recherche", "12. Limites et prochaines étapes",
        "13. Annexes techniques et protocole de captures",
    ]
    for item in items: doc.add_paragraph(clean(item), style="List Number")
    callout(doc, "Organisation du rapport", "Le document suit une progression narrative. Il part du cadrage initial, décrit les difficultés rencontrées, justifie les choix méthodologiques, présente les résultats observés puis distingue les acquis des travaux restant à consolider.")
    doc.add_page_break()


def add_dashboard(doc: Document) -> None:
    title(doc, "Tableau de bord de l’avancement")
    rows = [
        ("Chaîne complète", "Fonctionnelle", "Extraction du classeur, validation des métriques, puis production du rapport financier Word"),
        ("Phase 1", "Fonctionnelle à consolider", "Recherche hybride, unités, règles métier et validation"),
        ("Banque World Bank", "Active", "3 955 projets énergétiques classés"),
        ("Référentiel IRENA", "Actif", "1 687 observations sectorielles intégrées"),
        ("DuckDB", "Active", "29 427 observations et 29 143 événements de normalisation"),
        ("Phase 2", "Fonctionnelle à enrichir", "Calculs déterministes, benchmarks et rapport Word"),
        ("Tests", "En place", "85 tests automatisés au dernier décompte consolidé"),
        ("Évaluation multi classeurs", "À réaliser", "Condition nécessaire pour conclure sur la généralisation"),
    ]
    table = doc.add_table(rows=1, cols=3); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = "Table Grid"
    for i, h in enumerate(["Composant", "État", "Observation"]): set_cell_text(table.rows[0].cells[i], h, True, "FFFFFF", 9); shade(table.rows[0].cells[i], BLUE)
    for a, b, c in rows:
        cells = table.add_row().cells
        for i, val in enumerate([a, b, c]): set_cell_text(cells[i], val, i == 0, None, 9)
        if "À réaliser" in b or "consolider" in b or "enrichir" in b: shade(cells[1], "FFF2CC")
        else: shade(cells[1], "E2F0D9")
    doc.add_paragraph()
    p = doc.add_paragraph(clean("Ce tableau montre qu’une première chaîne complète est désormais opérationnelle, depuis la lecture du classeur jusqu’à la production du rapport financier. Les prochains travaux porteront principalement sur la mesure de la précision, l’uniformisation des unités et la vérification de la généralisation sur plusieurs modèles financiers.")); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_page_break()


def add_capture_protocol(doc: Document) -> None:
    title(doc, "Protocole de production des captures finales")
    rows = [
        ("Branche main", "Exécution réelle", "Extrait court montrant les erreurs 503 et l’absence de proposition"),
        ("Branche hybride publiée", "Exécution réelle", "Montage signalé montrant le repli déterministe et les temps longs"),
        ("Version locale actuelle", "Exécution réelle à produire", "Même classeur, même machine, état du cache et date indiqués"),
        ("Qualité avant", "Ancien commit ou trace réelle", "Même vérité terrain et mêmes métriques"),
        ("Qualité après", "Version actuelle", "Top 1, rappel à 5, rappel à 10 et rang réciproque moyen"),
        ("Validation", "Exécution réelle", "Valeur, unité, alternatives et navigation visibles"),
        ("Comparables", "Exécution réelle", "Vue synthétique puis décision individuelle"),
        ("DuckDB", "Requête réelle", "Tables et comptages, aucune donnée confidentielle"),
    ]
    table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Capture", "Nature", "Condition de validité"]): set_cell_text(table.rows[0].cells[i], h, True, "FFFFFF", 9); shade(table.rows[0].cells[i], BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row): set_cell_text(cells[i], val, i == 0, None, 8.5)
    doc.add_paragraph()
    callout(doc, "Règle de transparence", "Une sortie reconstruite peut illustrer un comportement ancien, mais elle ne doit jamais être présentée comme une mesure réelle. Les graphiques de performance ne devront utiliser que des exécutions reproductibles ou des traces historiques dont l’origine est conservée.", "FCE4D6")
    doc.add_page_break()


def add_execution_comparison(doc: Document) -> None:
    title(doc, "Comparaison des deux exécutions historiques")
    rows = [
        ("Libellés collectés", "4 192", "5 606"),
        ("Temps de collecte", "2,2 secondes", "2,7 secondes"),
        ("Ouverture du modèle", "0,5 seconde", "80,8 secondes"),
        ("Métriques traitées", "22", "22"),
        ("Routages LLM", "21", "22"),
        ("Erreurs 503 affichées", "23", "0"),
        ("Limites de quota affichées", "40", "22"),
        ("Métriques proposées", "1", "22"),
        ("Métriques laissées à explorer", "21", "0"),
        ("Somme des temps par métrique", "Non disponible", "1 024,3 secondes"),
        ("Durée minimale estimée", "Non calculable", "18 minutes et 28 secondes"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Indicateur", "origin/main", "origin/feature/architecture hybride"]):
        set_cell_text(table.rows[0].cells[i], h, True, "FFFFFF", 8.5)
        shade(table.rows[0].cells[i], BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, i == 0, None, 8.5)
        if row[0] in {"Métriques proposées", "Métriques laissées à explorer"}:
            shade(cells[1], "FCE4D6")
            shade(cells[2], "E2F0D9")
        if row[0] in {"Ouverture du modèle", "Somme des temps par métrique", "Durée minimale estimée"}:
            shade(cells[2], "FFF2CC")
    doc.add_paragraph()
    callout(doc, "Conclusion de la comparaison", "La branche hybride publiée marque un progrès net de résilience : elle conserve une proposition déterministe pour les vingt deux métriques malgré l’échec des appels Gemini. Elle ne constitue toutefois pas encore une solution performante, puisque le routage LLM demeure systématique et que l’étape nécessite au minimum dix huit minutes et vingt huit secondes.")
    doc.add_paragraph()
    callout(doc, "Précaution d’interprétation", "Les deux exécutions ont été réalisées sur des versions distinctes du code et dans un contexte de disponibilité dégradée de Gemini. Elles démontrent le comportement de repli de chaque architecture, mais ne suffisent pas à isoler statistiquement l’effet du réseau, du cache et de la machine.", "FCE4D6")
    doc.add_page_break()


def add_appendices(doc: Document) -> None:
    page(doc, "Annexe A : commandes principales du pipeline", [
        "L’exécution complète part du classeur Excel avec la commande python arsel_analyse.py suivie du chemin du modèle. Le système conduit alors la collecte, la recherche, la validation, la création du registre et l’analyse financière jusqu’au document Word.",
        "Lorsque le registre hypotheses_validees.json existe déjà, la commande python run_analysis.py hypotheses_validees.json permet de reprendre directement l’analyse sans rouvrir le classeur. La commande python run_phase2.py hypotheses_validees.json reste disponible pour exécuter uniquement la seconde phase.",
        "La banque de benchmarks possède ses propres commandes d’ingestion, de contrôle, de sélection des comparables et de promotion. Cette séparation évite qu’une mise à jour d’une source externe modifie silencieusement une analyse en cours.",
    ])
    page(doc, "Annexe B : métriques d’évaluation proposées", [
        "La précision en première position mesure la part des métriques pour lesquelles la bonne cellule est proposée directement. Le rappel à cinq et à dix mesure si la bonne cellule apparaît dans la liste courte, même lorsqu’elle n’est pas classée première. Le rang réciproque moyen mesure la qualité globale de l’ordre des candidats.",
        "L’évaluation doit aussi séparer la fidélité de la valeur et celle de l’unité. Une cellule correcte associée à une unité erronée ne peut pas être considérée comme un succès complet, car elle peut entraîner un calcul faux de plusieurs ordres de grandeur.",
        "Le temps total, le temps par métrique, le nombre d’appels au modèle, le taux d’erreur de service et le recours à la validation humaine doivent être enregistrés. Ces mesures permettront de comparer objectivement les architectures et les modèles.",
    ])
    page(doc, "Annexe C : glossaire", [
        "TF IDF désigne une méthode de pondération lexicale qui valorise les mots caractéristiques d’un libellé. Les embeddings sont des représentations numériques permettant de rapprocher des expressions sémantiquement voisines. DuckDB est une base analytique embarquée adaptée aux traitements tabulaires locaux.",
        "Le TRI est le taux de rentabilité interne. Le WACC est le coût moyen pondéré du capital. Le DSCR mesure la couverture du service de la dette. Le CAPEX représente les dépenses d’investissement et l’OPEX les dépenses d’exploitation. Le LCOE exprime le coût actualisé de l’énergie.",
        "Le staging est une base intermédiaire reconstruite et contrôlée avant promotion. Un snapshot immuable est une copie brute conservée sans écrasement. Le checksum SHA 256 sert à identifier exactement le contenu d’un fichier et à détecter toute modification.",
    ])
    page(doc, "Annexe D : éléments attendus pour la prochaine revue", [
        "La prochaine revue devrait présenter une comparaison chronométrée de l’ancienne et de la nouvelle architecture, une évaluation de la recherche sur plusieurs classeurs, un exemple complet de correction de valeur et d’unité, ainsi qu’un exemple de benchmark dont toutes les conversions sont justifiées.",
        "Elle devrait également montrer la vue d’ensemble des projets comparables, les raisons du classement, les projets finalement approuvés et l’effet de cette sélection sur les statistiques produites.",
        "Enfin, un exemple de conclusion enrichie par un modèle de langage pourra être présenté uniquement après stabilisation des calculs déterministes. Cette conclusion devra rester rattachée aux faits, aux écarts et aux sources fournis au modèle.",
    ])


def main() -> None:
    OUT.mkdir(exist_ok=True); ASSETS.mkdir(exist_ok=True)
    figures = {
        "architecture": ASSETS / "architecture_globale.png",
        "architecture_detailed": ASSETS / "architecture_globale_detaillee.png",
        "retrieval": ASSETS / "recherche_hybride.png",
        "bank": ASSETS / "benchmark_bank.png",
        "provenance": ASSETS / "provenance.png",
        "capture_main": ASSETS / "capture_01_main_erreurs_503.png",
        "capture_structures": ASSETS / "capture_02_structures_inconnues.png",
        "capture_units": ASSETS / "capture_03_unites_valeurs_aberrantes.png",
        "capture_availability": ASSETS / "capture_04_indisponibilite.png",
    }
    architecture_figure(figures["architecture"]); detailed_architecture_figure(figures["architecture_detailed"])
    retrieval_figure(figures["retrieval"])
    bank_figure(figures["bank"]); provenance_figure(figures["provenance"])

    missing_captures = [str(path) for key, path in figures.items() if key.startswith("capture_") and not path.exists()]
    if missing_captures:
        raise FileNotFoundError("Captures originales manquantes : " + ", ".join(missing_captures))

    doc = Document(); configure(doc); add_cover(doc); add_contents(doc); add_dashboard(doc)
    add_figure_page(doc, "Synoptique général du système", figures["architecture"], "Figure 1. Architecture fonctionnelle actuelle.", "Le modèle Excel demeure intact. La validation crée un registre indépendant. Les benchmarks sont gouvernés séparément et la Phase 2 ne travaille que sur des données validées et structurées.")
    add_architecture_chapter(doc, figures["architecture_detailed"])

    sections = parse_source()
    figure_insertions = {
        "3. Avancement de la première phase : extraction et validation": ("Recherche hybride des métriques", figures["retrieval"], "Figure 3. Fusion de preuves lexicales, sémantiques et structurelles."),
        "5. Constitution de la banque de benchmarks": ("Architecture de la banque de benchmarks", figures["bank"], "Figure 8. Cycle d’ingestion et de promotion contrôlée."),
        "7. Évolution de l’architecture par rapport au cadrage initial": ("Traçabilité de bout en bout", figures["provenance"], "Figure 9. Provenance depuis la cellule jusqu’à la conclusion."),
    }
    for sec_title, paras in sections:
        title(doc, sec_title)
        for paragraph_text in paras:
            p = doc.add_paragraph(clean(paragraph_text))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if sec_title in figure_insertions:
            h, img, cap = figure_insertions[sec_title]
            add_figure_page(doc, h, img, cap, "Le synoptique matérialise les responsabilités de chaque composant. Il sert aussi de support pour expliquer pourquoi le modèle de langage n’est plus placé au centre de toutes les décisions.")
        if sec_title == "3. Avancement de la première phase : extraction et validation":
            add_phase1_execution_captures(doc, figures)
            add_retrieval_theory(doc, figures)
            add_execution_comparison(doc)

    add_capture_protocol(doc)
    add_appendices(doc)

    while len(doc.paragraphs) < 10:
        doc.add_paragraph()
    doc.save(TARGET)
    print(json.dumps({"document": str(TARGET), "assets": str(ASSETS), "sections_source": len(sections)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
